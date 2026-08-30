from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from urllib.parse import unquote
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


SUPPLIER_HEADERS = [
    "Компания",
    "Сайт",
    "Телефоны",
    "Email",
    "Комментарий",
    "Реестр Минпромторга",
]
SPREADSHEETML_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
FONT_CHILD_ORDER = {
    name: index
    for index, name in enumerate(
        (
            "b",
            "i",
            "strike",
            "outline",
            "shadow",
            "condense",
            "extend",
            "sz",
            "u",
            "vertAlign",
            "color",
            "name",
            "charset",
            "family",
            "scheme",
        )
    )
}
COMMENT_LIMIT = 260
PROCUREMENT_REPORT_DISCLAIMER = (
    "Важно: отчёт подготовлен с помощью ИИ и предназначен для быстрой оценки закупочной документации. "
    "Критичные юридические, финансовые и технические условия сверяйте с официальными документами закупки. "
    "Отчёт не заменяет профессиональную проверку; решения по участию, цене и обязательствам принимает пользователь."
)
REGISTRY_FALLBACK_REPORT_DISCLAIMER = (
    "Важно: соответствие указанных поставщиков и товаров конкретным записям реестра Минпромторга "
    "не подтверждено. Поставщики найдены и проверены по открытым сайтам; перед закупкой запросите "
    "у них актуальные реестровые сведения."
)
QUOTE_REQUEST_INTRO = (
    "Просим выставить счёт или направить коммерческое предложение по указанным ниже товарам. "
    "В предложении просим указать цену, срок поставки, условия оплаты, документы качества и условия доставки."
)


MATCH_LEVEL_LABELS = {
    "exact": "точное совпадение",
    "adjacent": "смежная категория",
    "profile": "профильный поставщик",
    "reject": "не подтверждено",
}


def _save_xlsx(workbook: Workbook, path: Path) -> None:
    """Save an XLSX whose font nodes follow the OpenXML schema sequence."""
    workbook.save(path)
    _normalize_xlsx_font_order(path)


def _normalize_xlsx_font_order(path: Path) -> None:
    """Rewrite only styles.xml when openpyxl emitted non-canonical font ordering."""
    styles_part = "xl/styles.xml"
    font_tag = f"{{{SPREADSHEETML_NS}}}font"
    fonts_tag = f"{{{SPREADSHEETML_NS}}}fonts"

    with zipfile.ZipFile(path, "r") as source:
        try:
            styles = source.read(styles_part)
        except KeyError:
            return

        root = ET.fromstring(styles)
        changed = False
        for font in root.findall(f".//{fonts_tag}/{font_tag}"):
            children = list(font)
            ordered = sorted(
                children,
                key=lambda child: FONT_CHILD_ORDER.get(child.tag.rsplit("}", 1)[-1], len(FONT_CHILD_ORDER)),
            )
            if children != ordered:
                font[:] = ordered
                changed = True

        if not changed:
            return

        ET.register_namespace("", SPREADSHEETML_NS)
        normalized_styles = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        entries = [(entry, source.read(entry.filename)) for entry in source.infolist()]
        archive_comment = source.comment

    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f"{path.stem}-styles-",
        suffix=path.suffix,
        dir=path.parent,
    )
    os.close(descriptor)
    temporary_path = Path(temporary_name)
    try:
        with zipfile.ZipFile(temporary_path, "w") as target:
            target.comment = archive_comment
            for entry, content in entries:
                target.writestr(entry, normalized_styles if entry.filename == styles_part else content)
        os.replace(temporary_path, path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def _style_range(
    ws,
    row: int,
    start_col: int,
    end_col: int,
    *,
    fill: PatternFill | None = None,
    font: Font | None = None,
    align: Alignment | None = None,
    border: Border | None = None,
) -> None:
    if end_col > start_col:
        ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
    for c in range(start_col, end_col + 1):
        cell = ws.cell(row=row, column=c)
        if fill:
            cell.fill = fill
        if font:
            cell.font = font
        if align:
            cell.alignment = align
        if border:
            cell.border = border


def _product_fit_badge(row: dict) -> tuple[str, Font, PatternFill]:
    product_fit = str(row.get("product_fit") or "").strip().lower()
    if product_fit == "exact":
        return (
            "Точный товар",
            Font(name="Calibri", size=10, bold=True, color="15803D"),
            PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        )
    if product_fit == "analog":
        return (
            "Аналог",
            Font(name="Calibri", size=10, bold=True, color="B45309"),
            PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid"),
        )
    if product_fit == "category":
        return (
            "Категория",
            Font(name="Calibri", size=10, color="334155"),
            PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid"),
        )
    if product_fit == "profile":
        return (
            "Профиль компании",
            Font(name="Calibri", size=10, color="334155"),
            PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid"),
        )
    return (
        "Уточнить",
        Font(name="Calibri", size=10, color="64748B"),
        PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid"),
    )


def _registry_badge(row: dict) -> tuple[str, Font, PatternFill]:
    match = row.get("minprom_registry_match") if isinstance(row.get("minprom_registry_match"), dict) else {}
    if match.get("matched"):
        reg_no = _clean_comment_text(match.get("registry_number") or "")
        if not reg_no and match.get("evidence"):
            ev_m = re.search(r"(?:заключение|срок действия/заключение|реестровый номер|первичный)[:\s]*([A-Z0-9/-]+)", str(match.get("evidence") or ""), re.I)
            if ev_m:
                reg_no = ev_m.group(1).strip()
        label = f"№ {reg_no} (ГИСП)" if reg_no else "Подтверждён (ГИСП)"
        return (
            label,
            Font(name="Calibri", size=10, bold=True, color="15803D"),
            PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
        )
    status = str(row.get("minprom_registry_status") or "").strip().lower()
    origin = str(row.get("supplier_search_origin") or "").strip()
    policy = str(row.get("supplier_search_policy") or "").strip()
    if status == "empty" or origin == "ordinary_fallback" or policy in ("minprom_registry_only", "minprom_registry_priority"):
        return (
            "Обычный поиск",
            Font(name="Calibri", size=10, color="64748B"),
            PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid"),
        )
    return (
        "—",
        Font(name="Calibri", size=10, color="94A3B8"),
        PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid"),
    )


def _calc_supplier_row_h(cells: list[tuple[str, int]], line_h: int = 15, min_h: int = 24) -> int:
    max_lines = 1
    for val, col_w in cells:
        cleaned = str(val or "").strip()
        if not cleaned:
            continue
        chars_per_line = max(8, int(col_w * 0.85))
        for part in cleaned.split("\n"):
            lines = max(1, (len(part) + chars_per_line - 1) // chars_per_line)
            if lines > max_lines:
                max_lines = lines
    return max(min_h, max_lines * line_h + 8)


def _write_quote_request_sheet(wb: Workbook, *, title: str, subject: str = "") -> None:
    ws = wb.create_sheet(title="Запрос КП")
    ws.views.sheetView[0].showGridLines = True
    base_title = _clean_comment_text(subject) or _clean_comment_text(title) or "Продукция по ТЗ"

    # Шапка
    ws.append(["TenderLex | ГОТОВЫЙ ТЕКСТ ЗАПРОСА КП ДЛЯ КОПИРОВАНИЯ (НАЖМИТЕ НА ЯЧЕЙКУ A3 И СКОПИРУЙТЕ CTRL+C)"])
    _style_range(ws, 1, 1, 2, font=Font(name="Calibri", size=13, bold=True, color="0F172A"), align=Alignment(vertical="center"))
    ws.row_dimensions[1].height = 26

    ws.append([f"Предмет закупки / номенклатура: {base_title}"])
    _style_range(ws, 2, 1, 2, font=Font(name="Calibri", size=11, bold=True, color="334155"), align=Alignment(vertical="center"))
    ws.row_dimensions[2].height = 22

    # Скомпилированный текст в одной ячейке A3 для моментального копирования в Word / Email
    quote_text = (
        f"Тема письма: Запрос коммерческого предложения / счёта на поставку: {base_title}\n\n"
        "Добрый день, отдел продаж!\n\n"
        "Просим Вас предоставить коммерческое предложение (счёт) на поставку следующей продукции:\n"
        f"• {base_title} (согласно техническому заданию / спецификации заказчика)\n\n"
        "В коммерческом предложении просим обязательно указать:\n"
        "1. Стоимость за единицу и общую стоимость продукции (с учётом НДС).\n"
        "2. Фактическое наличие на складе и минимальные сроки отгрузки / производства.\n"
        "3. Условия и стоимость доставки до объекта либо возможность самовывоза.\n"
        "4. Наличие паспортов качества, сертификатов соответствия и реестровых записей ГИСП (при наличии).\n"
        "5. Срок действия коммерческого предложения и условия оплаты (аванс / постоплата).\n\n"
        "Ответ и коммерческое предложение просим направить в ответном письме либо по телефону.\n\n"
        "С уважением,\n"
        "Отдел закупок и снабжения"
    )

    ws.append([quote_text])
    c = ws.cell(row=3, column=1)
    c.font = Font(name="Calibri", size=11, color="0F172A")
    c.fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    c.border = Border(
        left=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"),
        bottom=Side(style="thin", color="CBD5E1"),
    )
    c.alignment = Alignment(vertical="top", horizontal="left", wrap_text=True)
    ws.row_dimensions[3].height = 360
    ws.column_dimensions["A"].width = 115


def write_supplier_xlsx(path: str | Path, rows: list[dict], *, title: str, target: int, subject: str = "", policy: str = "") -> Path:
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Поставщики"
    ws.views.sheetView[0].showGridLines = True

    # 1. Шапка документа
    brand_title = f"TenderLex | {_supplier_report_heading(title, subject)}"
    ws.append([brand_title])
    _style_range(ws, 1, 1, len(SUPPLIER_HEADERS), font=Font(name="Calibri", size=14, bold=True, color="0F172A"), align=Alignment(vertical="center"))
    ws.row_dimensions[1].height = 28

    # 2. Подзаголовок (ТЗ + Режим)
    policy_label = _supplier_policy_label(policy) or "Режим: Поиск поставщиков (Обычный)"
    item_title = _clean_comment_text(subject) or _clean_comment_text(title) or "Спецификация"
    ws.append([f"Предмет закупки / ТЗ: {item_title} | {policy_label}"])
    _style_range(ws, 2, 1, len(SUPPLIER_HEADERS), font=Font(name="Calibri", size=11, bold=True, color="334155"), align=Alignment(vertical="center"))
    ws.row_dimensions[2].height = 22

    # 3. Сводка / KPI (без лишней плашки "как работать")
    summary = _supplier_count_summary(rows, target)
    is_fallback = _is_registry_fallback_report(rows)
    if is_fallback:
        summary = f"{REGISTRY_FALLBACK_REPORT_DISCLAIMER}\n\n{summary}"
    ws.append([summary])
    summary_fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid") if is_fallback else PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    summary_font = Font(name="Calibri", size=10, bold=is_fallback, color="9C2A10" if is_fallback else "1E293B")
    _style_range(ws, 3, 1, len(SUPPLIER_HEADERS), fill=summary_fill, font=summary_font, align=Alignment(vertical="center", wrap_text=True))
    ws.row_dimensions[3].height = 36 if is_fallback else 24

    # 4. Разделитель
    ws.append([None] * len(SUPPLIER_HEADERS))
    ws.row_dimensions[4].height = 10

    # 5. Заголовки таблицы: Компания, Сайт, Телефоны, Email, Комментарий, Реестр Минпромторга
    ws.append(SUPPLIER_HEADERS)
    header_row = 5
    header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    for cell in ws[header_row]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")
    ws.row_dimensions[header_row].height = 26

    # 6. Данные
    thin_border = Border(
        left=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"),
        bottom=Side(style="thin", color="CBD5E1"),
    )

    data_start_row = header_row + 1
    for row_idx, row in enumerate(rows, start=data_start_row):
        is_even = (row_idx % 2 == 0)
        base_bg = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid") if is_even else PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
        
        company = str(row.get("company_name") or "").strip()
        reg_text, reg_font, reg_fill = _registry_badge(row)
        site_raw = str(row.get("site") or "").strip()
        site_display = unquote(site_raw) if site_raw else ""
        phone = str(row.get("phone") or "").strip()
        email = str(row.get("email") or "").strip()
        comment = _client_supplier_comment(row)

        ws.append([company, site_display, phone, email, comment, reg_text])

        # Col 1: Компания
        c1 = ws.cell(row=row_idx, column=1)
        c1.font = Font(name="Calibri", size=10, bold=True, color="0F172A")
        c1.fill = base_bg
        c1.border = thin_border
        c1.alignment = Alignment(wrap_text=True, vertical="top")

        # Col 2: Сайт
        c2 = ws.cell(row=row_idx, column=2)
        if site_raw:
            c2.hyperlink = site_raw
            c2.font = Font(name="Calibri", size=10, color="0284C7", underline="single")
        else:
            c2.font = Font(name="Calibri", size=10, color="64748B")
        c2.fill = base_bg
        c2.border = thin_border
        c2.alignment = Alignment(wrap_text=True, vertical="top")

        # Col 3: Телефоны
        c3 = ws.cell(row=row_idx, column=3)
        c3.font = Font(name="Calibri", size=10, color="1E293B")
        c3.fill = base_bg
        c3.border = thin_border
        c3.alignment = Alignment(wrap_text=True, vertical="top")

        # Col 4: Email
        c4 = ws.cell(row=row_idx, column=4)
        c4.font = Font(name="Calibri", size=10, color="0F766E")
        c4.fill = base_bg
        c4.border = thin_border
        c4.alignment = Alignment(wrap_text=True, vertical="top")

        # Col 5: Комментарий
        c5 = ws.cell(row=row_idx, column=5)
        c5.font = Font(name="Calibri", size=10, color="334155")
        c5.fill = base_bg
        c5.border = thin_border
        c5.alignment = Alignment(wrap_text=True, vertical="top")

        # Col 6: Реестр Минпромторга (последний столбец)
        c6 = ws.cell(row=row_idx, column=6)
        c6.font = reg_font
        c6.fill = reg_fill
        c6.border = thin_border
        c6.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

        ws.row_dimensions[row_idx].height = _calc_supplier_row_h([(company, 32), (comment, 60)], min_h=24)

    widths = [32, 35, 22, 26, 60, 26]
    for column, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(column)].width = width
    
    # Без закрепления областей - страница свободно прокручивается
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(len(SUPPLIER_HEADERS))}{max(header_row, ws.max_row)}"

    # Вторая вкладка: Запрос КП
    _write_quote_request_sheet(wb, title=title, subject=subject)

    _save_xlsx(wb, out)
    return out


def _match_level_label(value: object) -> str:
    raw = str(value or "").strip()
    return MATCH_LEVEL_LABELS.get(raw, raw)


def _supplier_count_summary(rows: list[dict], target: int) -> str:
    counts = {"exact": 0, "analog": 0, "category": 0, "profile": 0}
    for row in rows:
        product_fit = str(row.get("product_fit") or "").strip().lower()
        if product_fit in counts:
            counts[product_fit] += 1

    unclassified = max(0, len(rows) - sum(counts.values()))
    parts = [
        f"контактов с точным техническим совпадением: {counts['exact']}",
        f"возможных аналогов: {counts['analog']}",
        f"категорийных кандидатов: {counts['category'] + counts['profile']}",
    ]
    if unclassified:
        parts.append(f"требуют классификации: {unclassified}")
    return f"Проверены сайты и контакты. Кандидатов: {len(rows)}; {', '.join(parts)}."


def _is_registry_fallback_report(rows: list[dict]) -> bool:
    return any(
        str(row.get("supplier_search_policy") or "").strip() == "minprom_registry_only"
        and str(row.get("supplier_search_origin") or "").strip() == "ordinary_fallback"
        and not bool(
            row.get("minprom_registry_match", {}).get("matched")
            if isinstance(row.get("minprom_registry_match"), dict)
            else False
        )
        for row in rows
    )


def _client_supplier_comment(row: dict) -> str:
    product_fit = str(row.get("product_fit") or "").strip().lower()
    product = _clean_comment_text(row.get("product") or row.get("procurement_item") or "")
    raw_comment = _clean_comment_text(row.get("comments") or "").replace("ИИ", "Проверка").replace("AI", "Проверка")
    detail = _short_product_or_comment(product, raw_comment)
    registry_note = _supplier_registry_note(row)
    if product_fit == "exact":
        if detail:
            return _join_supplier_comment(f"Точное соответствие: {detail}.", registry_note)
        return _join_supplier_comment("Точное соответствие.", registry_note)
    if product_fit == "analog":
        if detail:
            return _join_supplier_comment(f"Возможный аналог: {detail}. Сверить характеристики.", registry_note)
        return _join_supplier_comment("Возможный аналог. Сверить характеристики.", registry_note)
    if product_fit == "category":
        if detail:
            return _join_supplier_comment(f"Категория совпадает: {detail}. Конкретный товар не подтвержден.", registry_note)
        return _join_supplier_comment("Категория совпадает. Конкретный товар не подтвержден.", registry_note)
    if product_fit == "profile":
        return _join_supplier_comment("Профиль компании подходит. Наличие товара уточнить.", registry_note)
    if detail:
        return _join_supplier_comment(f"Соответствие требует уточнения: {detail}.", registry_note)
    return _join_supplier_comment("Соответствие требует уточнения.", registry_note)


def _join_supplier_comment(base: str, registry_note: str) -> str:
    if not registry_note:
        return _truncate_comment(base, COMMENT_LIMIT)
    return _truncate_comment(f"{base} {registry_note}", 360)


def _supplier_registry_note(row: dict) -> str:
    policy = str(row.get("supplier_search_policy") or "").strip()
    required = bool(row.get("minprom_registry_required"))
    if policy not in {"minprom_registry_only", "minprom_registry_priority"} and not required:
        return ""
    match = row.get("minprom_registry_match") if isinstance(row.get("minprom_registry_match"), dict) else {}
    if match.get("matched"):
        registry_number = _clean_comment_text(match.get("registry_number") or "")
        manufacturer = _clean_comment_text(match.get("manufacturer") or "")
        if not registry_number and match.get("evidence"):
            ev_m = re.search(r"(?:заключение|срок действия/заключение|реестровый номер|первичный)[:\s]*([A-Z0-9/-]+)", str(match.get("evidence") or ""), re.I)
            if ev_m:
                registry_number = ev_m.group(1).strip()
        if registry_number and manufacturer:
            return f"Реестр ГИСП Минпромторга: запись № {registry_number}, производитель {manufacturer}."
        if registry_number:
            return f"Реестр ГИСП Минпромторга: запись № {registry_number}."
        if manufacturer:
            return f"Реестр ГИСП Минпромторга: подтверждён производитель {manufacturer}."
        return "Реестр: запись подтверждена."
    status = str(row.get("minprom_registry_status") or "").strip().lower()
    origin = str(row.get("supplier_search_origin") or "").strip()
    if status == "empty":
        if origin == "ordinary_fallback" or policy == "minprom_registry_priority":
            return "Реестр: релевантная запись не найдена; поставщик найден обычным поиском."
        return "Реестр: релевантная запись не найдена."
    if status == "error":
        return "Реестр: проверка не выполнена."
    if status == "ok" and origin == "ordinary_fallback":
        return "Реестр: соответствие поставщика конкретной записи не подтверждено; поставщик найден обычным поиском."
    return ""


POLICY_DISPLAY_NAMES = {
    "minprom_registry_priority": "Режим: Поиск поставщиков (Реестр в приоритете)",
    "minprom_registry_only": "Режим: Поиск поставщиков (Только реестр)",
    "normal": "Режим: Поиск поставщиков (Обычный)",
}

def _supplier_report_heading(title: str, subject: str = "", policy: str = "") -> str:
    source = _clean_comment_text(title)
    item = _clean_comment_text(subject)
    base_title = item or source or "ТЗ"
    return f"Отчёт по ТЗ: {base_title}"


def _supplier_policy_label(policy: str) -> str:
    return POLICY_DISPLAY_NAMES.get(policy, "")


def _clean_comment_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip(" .,:;")


def _short_product_or_comment(product: str, comment: str) -> str:
    if product:
        return _truncate_comment(product, 130).rstrip(".")
    if not comment:
        return ""
    softened = _remove_registry_comment_fragments(_soften_supplier_claims(comment))
    if not softened:
        return ""
    first_sentence = re.split(r"(?<=[.!?])\s+", softened, maxsplit=1)[0]
    return _truncate_comment(first_sentence, 130).rstrip(".")


def _soften_supplier_claims(comment: str) -> str:
    value = str(comment or "")
    value = re.sub(r"что\s+полностью\s+соответствует", "что релевантно", value, flags=re.I)
    value = re.sub(r"профиль\s+полностью\s+соответствует", "профиль релевантен", value, flags=re.I)
    value = re.sub(r"полностью\s+соответствует\s+ТЗ", "может быть релевантно ТЗ", value, flags=re.I)
    value = re.sub(r"полностью\s+соответствует", "релевантно", value, flags=re.I)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def _remove_registry_comment_fragments(comment: str) -> str:
    value = re.sub(r"\s+", " ", str(comment or "")).strip()
    if not value:
        return ""
    parts = re.split(r"(?<=[.!?])\s+", value)
    kept = [
        part.strip()
        for part in parts
        if part.strip() and not re.search(r"(?:минпромторг\w*|гисп|реестр\w*|реестров\w*)", part, flags=re.I)
    ]
    return " ".join(kept)


def _truncate_comment(comment: str, limit: int = 500) -> str:
    value = str(comment or "").strip()
    if len(value) <= limit:
        return value
    value = value[:limit].rsplit(" ", 1)[0].rstrip(" .,:;")
    return f"{value}."


def _compact_comment_detail(comment: str, limit: int = 260) -> str:
    value = re.sub(r"\s+", " ", str(comment or "")).strip()
    if len(value) <= limit:
        return value
    selected = ""
    for sentence in re.split(r"(?<=[.!?])\s+", value):
        sentence = sentence.strip()
        if not sentence:
            continue
        candidate = f"{selected} {sentence}".strip()
        if len(candidate) > limit:
            break
        selected = candidate
        if len(selected) >= 160:
            break
    return selected or _truncate_comment(value, limit=limit)


def write_procurement_docx(path: str | Path, markdown: str, *, title: str) -> Path:
    return _write_markdown_docx(
        path,
        markdown,
        title=title or "Отчёт анализа закупки",
        intro=PROCUREMENT_REPORT_DISCLAIMER,
        intro_italic=True,
    )


def write_quote_request_docx(path: str | Path, markdown: str, *, title: str = "Запрос КП") -> Path:
    return _write_markdown_docx(path, markdown, title=title or "Запрос КП")


def _write_markdown_docx(
    path: str | Path,
    markdown: str,
    *,
    title: str,
    intro: str = "",
    intro_italic: bool = False,
) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    brand_emerald = RGBColor(4, 120, 87)       # #047857
    dark_emerald = RGBColor(6, 78, 59)        # #064E3B
    teal_emerald = RGBColor(15, 118, 110)     # #0F766E
    text_dark = RGBColor(15, 23, 42)

    # Top brand header
    h_top = doc.add_paragraph()
    h_top.paragraph_format.space_before = Pt(0)
    h_top.paragraph_format.space_after = Pt(2)
    r_b = h_top.add_run("TenderLex")
    r_b.font.bold = True
    r_b.font.size = Pt(14)
    r_b.font.color.rgb = brand_emerald

    r_p = h_top.add_run(" | ")
    r_p.font.size = Pt(14)
    r_p.font.color.rgb = RGBColor(148, 163, 184)

    r_t = h_top.add_run(title)
    r_t.font.bold = True
    r_t.font.size = Pt(13)
    r_t.font.color.rgb = text_dark

    if intro:
        intro_paragraph = doc.add_paragraph()
        intro_paragraph.paragraph_format.space_before = Pt(2)
        intro_paragraph.paragraph_format.space_after = Pt(6)
        intro_run = intro_paragraph.add_run(intro)
        intro_run.italic = intro_italic
        intro_run.font.size = Pt(8.5)
        intro_run.font.color.rgb = RGBColor(100, 116, 139)

    lines = _remove_okpd_codes(str(markdown or "")).splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if _is_markdown_table_start(lines, index):
            index = _add_markdown_table(doc, lines, index)
            continue
        if re.fullmatch(r"-{3,}", line):
            doc.add_paragraph("")
            index += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[2:])
            r.font.bold = True
            r.font.size = Pt(13.5)
            r.font.color.rgb = brand_emerald
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line[3:])
            r.font.bold = True
            r.font.size = Pt(12.5)
            r.font.color.rgb = dark_emerald
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(line[4:])
            r.font.bold = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = dark_emerald
        elif line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line[5:])
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = dark_emerald
        elif line.startswith(("- ", "* ", "1. ", "2. ", "3. ", "4. ", "5. ")):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(2)
            _add_markdown_runs(paragraph, line, font_size=Pt(9.5), font_color=text_dark)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(2)
            _add_markdown_runs(paragraph, line, font_size=Pt(9.5), font_color=text_dark)
        index += 1
    doc.save(out)
    return out


def _is_markdown_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    separator = lines[index + 1].strip()
    return current.startswith("|") and current.endswith("|") and bool(re.fullmatch(r"\|[\s:\-|]+\|", separator))


def _parse_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [
        cell.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
        for cell in cells
    ]


def _add_markdown_table(doc, lines: list[str], index: int) -> int:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Inches, Pt, RGBColor

    rows: list[list[str]] = [_parse_table_row(lines[index])]
    index += 2
    while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
        rows.append(_parse_table_row(lines[index]))
        index += 1
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set table width = 6.97 inches (10037 dxa)
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="10037" w:type="dxa"/>')
    tblPr.append(tblW)

    _apply_table_column_widths(table, _table_column_widths(rows[0], width))

    # Set full grid borders
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Repeat header & prevent row split
    for row in table.rows[:1]:
        row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    for row in table.rows:
        row._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for row_index, row in enumerate(rows):
        fill_color = "064E3B" if row_index == 0 else ("F4FBF7" if row_index % 2 == 1 else "FFFFFF")
        for col_index in range(width):
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
            cell._tc.get_or_add_tcPr().append(shd)

            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            value = _remove_okpd_codes(row[col_index] if col_index < len(row) else "")
            _add_markdown_runs(paragraph, value)

            if row_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                if col_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return index


def _table_column_widths(headers: list[str], width: int) -> list[int] | None:
    normalized = [_normalize_table_header(header) for header in headers]
    if normalized[:5] == ["№", "наименование", "характеристики", "ед.изм.", "кол-во"]:
        if width >= 6 and normalized[5] == "примечание":
            return [520, 2100, 4700, 850, 850, 1340]
        return [520, 2300, 5600, 850, 850]
    return None


def _normalize_table_header(value: object) -> str:
    normalized = re.sub(r"\s+", " ", str(value or "").strip().lower())
    normalized = normalized.replace("ед. изм.", "ед.изм.")
    normalized = normalized.replace("ед изм", "ед.изм.")
    normalized = normalized.replace("количество", "кол-во")
    return normalized


def _apply_table_column_widths(table, widths: list[int] | None) -> None:
    if not widths:
        return
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Twips

    table.autofit = False
    for col_index, width in enumerate(widths):
        if col_index >= len(table.columns):
            break
        table.columns[col_index].width = Twips(width)
    for row in table.rows:
        for col_index, cell in enumerate(row.cells):
            if col_index >= len(widths):
                continue
            cell.width = Twips(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0


def _add_markdown_runs(
    paragraph,
    text: str,
    *,
    font_size=None,
    font_color=None,
) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", str(text or ""))
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        for line_index, segment in enumerate(value.split("\n")):
            if line_index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(segment)
            run.bold = bold
            if font_size is not None:
                run.font.size = font_size
            if font_color is not None:
                run.font.color.rgb = font_color


def _remove_okpd_codes(text: str) -> str:
    value = str(text or "")
    okpd_name = r"(?:ОКПД\s*2?|OKPD\s*2?|КТРУ|KTRU)"
    value = re.sub(
        rf"\s*[\(\[]\s*(?:код\s+)?{okpd_name}\s*[:№#Nn\-–—]?\s*[\d.\s/-]+[\)\]]",
        "",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        rf"(?:код\s+)?{okpd_name}\s*[:№#Nn\-–—]?\s*\d+(?:\.\d+){{1,6}}\b",
        "",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(r"\b\d{2}\.\d{2}\.\d{2}(?:\.\d{1,3}){0,3}\b", "", value)
    value = re.sub(r"[ \t]{2,}", " ", value)
    value = re.sub(r"\s+([,.;:])", r"\1", value)
    value = re.sub(r"\(\s*\)", "", value)
    value = re.sub(r"\[\s*\]", "", value)
    value = re.sub(r"\n[ \t]+", "\n", value)
    return value.strip()


def write_evidence(path: str | Path, payload: dict) -> Path:
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def zip_paths(zip_path: str | Path, files: list[Path]) -> Path:
    out = Path(zip_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_path in files:
            archive.write(file_path, arcname=file_path.name)
    return out

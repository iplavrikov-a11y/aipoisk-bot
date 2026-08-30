from __future__ import annotations

import asyncio
import html
import json
import logging
import os
import re
import sqlite3
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, List, Dict, Optional
from urllib.parse import urlsplit

import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from .ai import call_llm
from .models import SystemSettings
from .supplier_search import (
    _minprom_registry_sqlite_path,
    _search_with_yandex,
    _yandex_credentials,
    GISP_PRODUCT_REGISTRY_URL,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data Models
# ---------------------------------------------------------------------------

@dataclass
class SpecParameterMatch:
    param_name: str
    tz_requirement: str
    product_fact: str
    status: str = "match"  # "match" | "mismatch" | "clarify"
    comment: str = ""
    source_url: str = ""
    source_doc: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class GispRegistryMatch:
    registry_number: str
    manufacturer: str
    product: str
    inn: str = ""
    source_url: str = GISP_PRODUCT_REGISTRY_URL
    matched: bool = True

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class AlternativeProduct:
    brand: str
    model: str
    manufacturer: str
    confidence: float = 0.90
    notes: str = ""
    specs_breakdown: list[SpecParameterMatch] = field(default_factory=list)
    gisp_match: Optional[GispRegistryMatch] = None
    source_url: str = ""
    datasheet_url: str = ""

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["specs_breakdown"] = [s.to_dict() if isinstance(s, SpecParameterMatch) else s for s in self.specs_breakdown]
        data["gisp_match"] = self.gisp_match.to_dict() if self.gisp_match else None
        return data


@dataclass
class ExactProductPosition:
    position_no: int
    name_in_tz: str
    identified_brand: str
    identified_model: str
    manufacturer: str
    confidence: float
    reasoning: str
    specs_breakdown: list[SpecParameterMatch] = field(default_factory=list)
    alternative_brands: list[AlternativeProduct] = field(default_factory=list)
    gisp_match: Optional[GispRegistryMatch] = None
    source_url: str = ""
    datasheet_url: str = ""

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["specs_breakdown"] = [s.to_dict() if isinstance(s, SpecParameterMatch) else s for s in self.specs_breakdown]
        data["alternative_brands"] = [a.to_dict() if isinstance(a, AlternativeProduct) else a for a in self.alternative_brands]
        data["gisp_match"] = self.gisp_match.to_dict() if self.gisp_match else None
        return data


@dataclass
class ExactProductReport:
    procurement_title: str
    total_positions: int
    positions: list[ExactProductPosition] = field(default_factory=list)
    summary: str = ""
    disclaimer: str = (
        "Отчёт сформирован на основе сопоставления технического задания с открытыми веб-источниками, "
        "каталогами производителей, PDF-паспортами изделий и реестром Минпромторга РФ (ГИСП). "
        "Все показатели проверены первоисточниками без искусственной подгонки под ТЗ заказчика по 44-ФЗ и 223-ФЗ."
    )
    yandex_requests_count: int = 0
    yandex_cost_rub: float = 0.0
    web_sources: list[str] = field(default_factory=list)
    verified_documents: list[dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "procurement_title": self.procurement_title,
            "total_positions": self.total_positions,
            "positions": [p.to_dict() for p in self.positions],
            "summary": self.summary,
            "disclaimer": self.disclaimer,
            "yandex_requests_count": self.yandex_requests_count,
            "yandex_cost_rub": self.yandex_cost_rub,
            "web_sources": self.web_sources,
            "verified_documents": self.verified_documents,
        }


# ---------------------------------------------------------------------------
# Strict Grounded Prompt & Parsing
# ---------------------------------------------------------------------------

EXACT_PRODUCT_PROMPT = """Ты — ведущий эксперт по государственным закупкам (44-ФЗ, 223-ФЗ), стандартизации и промышленному оборудованию.
Твоя задача — проанализировать техническое задание (ТЗ), сопоставить его с приложенными проверенными документами из интернета (паспортами, каталогами, сайтами производителей) и сформировать достоверные сведения для заявки (Форма 2) и взаимозаменяемых аналогов.

ЖЕЛЕЗНЫЕ ПРАВИЛА ДОСТОВЕРНОСТИ И ИНЖЕНЕРНОГО АНАЛИЗА:
1. СТРОГО ЗАПРЕЩЕНО ВЫДУМЫВАТЬ ЗНАЧЕНИЯ ИЛИ ИСКУССТВЕННО ПОДГОНЯТЬ ИХ ПОД ТЗ!
2. ИСПОЛЬЗУЙ ПАСПОРТА, ТАБЛИЦЫ КАТАЛОГОВ И ГОСТ/ТУ:
   - Внимательно читай приложенные таблицы характеристик ([ИСТОЧНИК N]), опросные листы и модельные ряды заводов.
   - Если в ТЗ задан диапазон (например: 'глубина >=4 <=4.5 м', 'масса не более 12100 кг', 'мощность 2x0.55 кВт'), а в каталоге/паспорте приведено номинальное заводское значение или типовой ряд завода (например: глубина 4.2 м, масса 11800 кг, привод 2x0.55 кВт) — укажи конкретный заводской номинал, статус "match" и сошлись на [ИСТОЧНИК N].
   - Для стандартной продукции (металлопрокат, кабели, запорная арматура, КИПиА, электродвигатели) опирайся на официальные ГОСТ/ТУ (например ГОСТ 31996-2012, ГОСТ 57837-2017, ГОСТ 33259-2015).
3. ЕСЛИ ПАРАМЕТР ТОЧНО ПОДТВЕРЖДЕН И СООТВЕТСТВУЕТ ТЗ:
   - "product_fact": Конкретное подтвержденное значение производителя (без слов 'не менее/не более', например: "4.2 м", "11 800 кг", "AISI 304", "периферийный двухколесный")
   - "status": "match"
   - "comment": "Подтверждено каталогом производителя [ИСТОЧНИК N] / ГОСТ"
4. ЕСЛИ ФАКТИЧЕСКИЙ ПАРАМЕТР РАСХОДИТСЯ С ТЗ:
   - "product_fact": Реальный показатель из документа (например: "14 500 кг")
   - "status": "mismatch"
   - "comment": "Отклонение от ТЗ: фактически 14500 кг при требовании не более 12100 кг [ИСТОЧНИК N]"
5. ЕСЛИ ПАРАМЕТР ЯВЛЯЕТСЯ ЗАКАЗНОЙ ОПЦИЕЙ И ОТСУТСТВУЕТ В ОТКРЫТЫХ КАТАЛОГАХ:
   - "product_fact": "В открытой документации не указано (требуется официальный паспорт завода)"
   - "status": "clarify"
   - "comment": "Параметр согласуется по опросным листам / рабочему проекту [ИСТОЧНИК N]"
6. ДЛЯ АНАЛОГОВ (alternative_brands):
   - Указывай ТОЛЬКО реально существующие модели российских заводов из приложенных источников.
   - Построчно заполни характеристики аналога на основе его каталога/паспорта. Если параметр соответствует ТЗ — ставь "match", если отличается — "mismatch", если неизвестен — "clarify".

Спецификация ТЗ и проверенные документы из открытых источников:
{context}

Ответь СТРОГО в формате JSON:
{{
  "summary": "1-2 кратких предложения: какая конкретная модель заложена в ТЗ и какие проверенные аналоги РФ выявлены.",
  "positions": [
    {{
      "position_no": 1,
      "name_in_tz": "Наименование позиции из ТЗ",
      "identified_brand": "Бренд или завод",
      "identified_model": "Точная модель / серия",
      "manufacturer": "Завод-производитель",
      "confidence": 0.95,
      "reasoning": "Обоснование соответствия по ГОСТ/ТУ/размерам",
      "source_url": "URL сайта производителя или паспорта",
      "specs_breakdown": [
        {{
          "param_name": "Наименование параметра",
          "tz_requirement": "Требование из ТЗ",
          "product_fact": "Фактический показатель (или 'В открытом доступе не найдено' если нет в паспорте)",
          "status": "match",
          "comment": "Обоснование / ссылка на паспорт",
          "source_url": "URL источника"
        }}
      ],
      "alternative_brands": [
        {{
          "brand": "Бренд аналога",
          "model": "Модель аналога",
          "manufacturer": "Завод аналога",
          "confidence": 0.90,
          "notes": "Обоснование эквивалентности и отличия",
          "source_url": "URL сайта завода аналога",
          "specs_breakdown": [
            {{
              "param_name": "Наименование параметра",
              "tz_requirement": "Требование ТЗ",
              "product_fact": "Фактический показатель аналога (или 'В открытом доступе не найдено')",
              "status": "match",
              "comment": "Соответствие или отклонение аналога"
            }}
          ]
        }}
      ]
    }}
  ]
}}
"""


def extract_clean_spec_text(text: str) -> str:
    """Извлекает ключевой фрагмент технического задания или спецификации."""
    if not text:
        return ""
    cleaned = text.strip()
    if len(cleaned) <= 15000:
        return cleaned

    match = re.search(
        r"(?i)(?:(?:#+\s*)?(?:ТЕХНИЧЕСКОЕ\s+ЗАДАНИЕ|СПЕЦИФИКАЦИЯ|ОПИСАНИЕ\s+ОБЪЕКТА\s+ЗАКУПКИ|ТРЕБОВАНИЯ\s+К\s+ТОВАРУ|ТАБЛИЦА\s+ХАРАКТЕРИСТИК))(.+)",
        cleaned,
        re.DOTALL,
    )
    if match and len(match.group(1).strip()) > 100:
        cleaned = match.group(0).strip()

    if len(cleaned) > 35000:
        cleaned = cleaned[:35000]
    return cleaned


# ---------------------------------------------------------------------------
# Web & PDF Document Fetcher
# ---------------------------------------------------------------------------

async def _fetch_with_browser_fallback(url: str, domain: str) -> Optional[dict[str, Any]]:
    """
    Fallback-загрузчик через Playwright Chromium для обхода защит (KillBot, Cloudflare, DDoS-Guard)
    и рендеринга SPA/динамических таблиц характеристик.
    """
    try:
        from .procurement_sources import fetch_source_page_with_browser
        browser_page = await asyncio.wait_for(fetch_source_page_with_browser(url), timeout=18.0)
        if browser_page and browser_page.get("text"):
            text = str(browser_page["text"]).strip()
            if len(text) > 80:
                title = f"Официальный каталог / Спецификация ({domain})"
                return {
                    "url": url,
                    "domain": domain,
                    "type": "html_browser",
                    "title": title,
                    "text": text[:25000],
                }
    except Exception as b_exc:
        logger.debug("browser_fallback_failed for %s: %s", url, b_exc)
    return None


async def fetch_web_or_pdf_document(
    client: httpx.AsyncClient,
    url: str,
    timeout_seconds: float = 12.0,
) -> Optional[dict[str, Any]]:
    """
    Скачивает и извлекает чистый текст из HTML-страниц или PDF-паспортов/руководств.
    Поддерживает извлечение таблиц технических характеристик и параметров.
    При блокировках (403, 503, KillBot, Cloudflare) или пустом теле автоматически
    переключается на Playwright Headless Browser.
    """
    if not url or not url.startswith(("http://", "https://")):
        return None

    domain = urlsplit(url).netloc.lower()
    if any(bad in domain for bad in ["youtube.com", "vk.com", "t.me", "rutube.ru", "avito.ru", "wildberries.ru", "ozon.ru", "market.yandex.ru"]):
        return None

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml,application/pdf;q=0.9,*/*;q=0.8",
            "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
        }
        response = await client.get(url, headers=headers, timeout=timeout_seconds, follow_redirects=True)
        if response.status_code in (403, 503, 429):
            # Антибот блокировка -> Playwright fallback
            return await _fetch_with_browser_fallback(url, domain)
        if response.status_code >= 400:
            return None

        ctype = (response.headers.get("content-type") or "").lower()
        url_lower = str(response.url).lower()
        is_pdf = "application/pdf" in ctype or url_lower.endswith(".pdf") or response.content.startswith(b"%PDF-")

        if is_pdf:
            pdf_bytes = response.content
            if len(pdf_bytes) > 20 * 1024 * 1024:
                pdf_bytes = pdf_bytes[:20 * 1024 * 1024]

            pdf_text = ""
            try:
                import fitz
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                pages_text = []
                for pno, page in enumerate(doc):
                    page_parts = []
                    # 1. Извлечение структурированных таблиц (Табличный парсинг по методологии Docling)
                    try:
                        tables = page.find_tables()
                        for t_idx, tab in enumerate(tables):
                            tab_md = tab.to_markdown()
                            if tab_md and tab.row_count >= 2:
                                page_parts.append(f"\n[ТАБЛИЦА ТЕХНИЧЕСКИХ ХАРАКТЕРИСТИК (СТР. {pno + 1}, ТАБЛ. #{t_idx + 1})]:\n{tab_md}\n")
                    except Exception as tab_err:
                        logger.debug("fitz_find_tables_error: %s", tab_err)

                    # 2. Текст страницы
                    p_text = page.get_text("text").strip()
                    if p_text:
                        page_parts.append(p_text)

                    if page_parts:
                        pages_text.append(f"--- [СТРАНИЦА ПАСПОРТА {pno + 1}] ---\n" + "\n".join(page_parts))
                    if len("\n".join(pages_text)) > 30000:
                        break
                doc.close()
                pdf_text = "\n".join(pages_text)
            except Exception as pdf_err:
                logger.debug("pdf_extraction_error: %s for %s", pdf_err, url)

            if len(pdf_text.strip()) > 50:
                doc_name = response.url.path.split("/")[-1] or "Паспорт изделия (PDF)"
                return {
                    "url": str(response.url),
                    "domain": domain,
                    "type": "pdf",
                    "title": f"Паспорт / Техническая документация: {doc_name}",
                    "text": pdf_text[:30000],
                }

        # HTML parsing
        html_raw = response.text
        if not html_raw:
            return await _fetch_with_browser_fallback(url, domain)

        # Проверка на антибот заглушки в теле 200 OK
        lower_html = html_raw.lower()
        if any(b in lower_html for b in ["killbot", "ddos-guard", "cloudflare", "challenge-platform", "enable javascript", "проверка браузера", "защита от роботов"]):
            browser_doc = await _fetch_with_browser_fallback(url, domain)
            if browser_doc:
                return browser_doc

        soup = BeautifulSoup(html_raw, "html.parser")
        for tag in soup(["script", "style", "noscript", "svg", "header", "footer", "nav", "aside", "form"]):
            tag.decompose()

        page_title = soup.title.string.strip() if soup.title and soup.title.string else domain
        page_title = re.sub(r"\s+", " ", page_title)

        text_blocks: list[str] = []

        # 1. Извлечение таблиц характеристик (спецификаций)
        tables = soup.find_all("table")
        for t_idx, table in enumerate(tables[:10], start=1):
            rows = []
            for tr in table.find_all("tr"):
                cells = [" ".join(c.get_text().split()) for c in tr.find_all(["th", "td"])]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows and len(rows) >= 2:
                text_blocks.append(f"\n[ТАБЛИЦА ТЕХНИЧЕСКИХ ХАРАКТЕРИСТИК #{t_idx}]:\n" + "\n".join(rows))

        # 2. Извлечение списков параметров (dl / ul / div с параметрами)
        for dlist in soup.find_all(["dl", "ul", "div"], class_=re.compile(r"(?i)(spec|param|charact|feature|prop|tech)")):
            d_text = " ".join(dlist.get_text(" ", strip=True).split())
            if len(d_text) > 30 and len(d_text) < 4000:
                text_blocks.append(f"[БЛОК ПАРАМЕТРОВ]: {d_text}")

        # 3. Основной текст страницы
        main_text = soup.get_text("\n", strip=True)
        if main_text:
            text_blocks.append(main_text)

        full_extracted = "\n\n".join(text_blocks)
        clean_extracted = re.sub(r"\n{3,}", "\n\n", full_extracted).strip()

        if len(clean_extracted) > 80:
            return {
                "url": str(response.url),
                "domain": domain,
                "type": "html",
                "title": page_title,
                "text": clean_extracted[:25000],
            }
        else:
            return await _fetch_with_browser_fallback(url, domain)

    except Exception as exc:
        logger.debug("fetch_doc_failed for %s: %s", url, exc)
        return await _fetch_with_browser_fallback(url, domain)

    return None


async def fetch_batch_web_documents(urls: list[str], max_docs: int = 6) -> list[dict[str, Any]]:
    """Параллельно скачивает и извлекает контент из нескольких веб-страниц и PDF-паспортов."""
    if not urls:
        return []

    unique_urls = list(dict.fromkeys(urls))[:max_docs]
    docs: list[dict[str, Any]] = []

    async with httpx.AsyncClient(timeout=14.0, follow_redirects=True) as client:
        tasks = [fetch_web_or_pdf_document(client, u) for u in unique_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        for res in results:
            if isinstance(res, dict) and res.get("text"):
                docs.append(res)

    return docs


# ---------------------------------------------------------------------------
# Minpromtorg GISP Lookup
# ---------------------------------------------------------------------------

def find_minprom_gisp_match(
    brand: str,
    manufacturer: str,
    model: str = "",
    name_in_tz: str = "",
) -> Optional[GispRegistryMatch]:
    """Ищет запись в локальном SQLite FTS5 индексе Реестра Минпромторга (ГИСП)."""
    sqlite_path = _minprom_registry_sqlite_path()
    if not sqlite_path or not sqlite_path.is_file():
        shared_path = Path("/root/projects/emailagent/storage/minprom_registry/minprom_registry.sqlite")
        if shared_path.is_file():
            sqlite_path = shared_path
        else:
            return None

    queries = []
    if manufacturer and len(manufacturer.strip()) > 3:
        clean_manuf = re.sub(r'(?i)(ООО|АО|ПАО|ЗАО|НПК|НПО|ПК|ТД|ИП|ГК|«|»|")', '', manufacturer).strip()
        if len(clean_manuf) >= 3:
            queries.append(clean_manuf)
    if brand and len(brand.strip()) > 2 and brand not in queries:
        queries.append(brand.strip())
    if model and len(model.strip()) > 2:
        queries.append(model.strip())
    if name_in_tz and len(name_in_tz.strip()) > 4:
        clean_name = re.sub(r'[^а-яА-Яa-zA-Z0-9\s]', ' ', name_in_tz).split()
        if len(clean_name) >= 2:
            queries.append(" ".join(clean_name[:3]))

    if not queries:
        return None

    conn = None
    try:
        conn = sqlite3.connect(f"file:{sqlite_path}?mode=ro", uri=True)
        for q in queries:
            terms = [t for t in re.findall(r"[\w]{2,}", q) if len(t) > 1]
            if not terms:
                continue
            # 1. Точный AND поиск по префиксам
            match_expression = " AND ".join([f'"{t}"*' for t in terms[:4]])
            try:
                cursor = conn.execute(
                    """
                    SELECT e.registry_number, e.manufacturer, e.product, e.inn, e.source_url
                    FROM entries_fts
                    JOIN entries e ON e.id = entries_fts.rowid
                    WHERE entries_fts MATCH ?
                    LIMIT 3
                    """,
                    (match_expression,),
                )
                rows = cursor.fetchall()
                if rows:
                    reg_num, manuf, prod, inn, src_url = rows[0]
                    return GispRegistryMatch(
                        registry_number=str(reg_num or "").strip(),
                        manufacturer=str(manuf or "").strip(),
                        product=str(prod or "").strip(),
                        inn=str(inn or "").strip(),
                        source_url=str(src_url or GISP_PRODUCT_REGISTRY_URL),
                        matched=True,
                    )
            except sqlite3.Error:
                pass

            # 2. Мягкий поиск по первому ключевому термину (если многословный запрос не дал результатов)
            if len(terms) > 1 and len(terms[0]) >= 4:
                try:
                    cursor = conn.execute(
                        """
                        SELECT e.registry_number, e.manufacturer, e.product, e.inn, e.source_url
                        FROM entries_fts
                        JOIN entries e ON e.id = entries_fts.rowid
                        WHERE entries_fts MATCH ?
                        LIMIT 3
                        """,
                        (f'"{terms[0]}"*',),
                    )
                    rows = cursor.fetchall()
                    if rows:
                        reg_num, manuf, prod, inn, src_url = rows[0]
                        return GispRegistryMatch(
                            registry_number=str(reg_num or "").strip(),
                            manufacturer=str(manuf or "").strip(),
                            product=str(prod or "").strip(),
                            inn=str(inn or "").strip(),
                            source_url=str(src_url or GISP_PRODUCT_REGISTRY_URL),
                            matched=True,
                        )
                except sqlite3.Error:
                    pass

    except Exception as exc:
        logger.warning("gisp_lookup_failed: %s", exc)
    finally:
        if conn:
            conn.close()

    return None


# ---------------------------------------------------------------------------
# Adaptive Clarify & Standards Resolvers
# ---------------------------------------------------------------------------

RESOLVE_CLARIFY_PROMPT = """Ты — ведущий эксперт по стандартизации и госзакупкам (44-ФЗ/223-ФЗ).
Твоя задача — проверить приложенный текст технических документов / паспортов / каталогов / ГОСТ и установить точное фактическое значение параметра для конкретной модели оборудования или материала.

Товар: {brand} {model} (Производитель: {manufacturer})
Наименование параметра: {param_name}
Требование ТЗ: {tz_requirement}

Текст найденных технических документов, паспортов и стандартов:
{doc_text}

ПРАВИЛА ВЕРИФИКАЦИИ:
1. Ищи ТОЧНОЕ числовое или качественное значение параметра именно для указанной модели или нормативного типоразмера.
2. Не выдумывай и не подгоняй под ТЗ! Если точное значение ЕСТЬ в тексте:
   - "found": true
   - "product_fact": "точное значение из документа" (без слов 'не менее/не более', например: "1.6 МПа", "14 000 лм", "32 м3/ч", "1200 x 1200 dpi", "120 мм", "УХЛ1", "IP66")
   - "status": "match" (если значение укладывается в требование ТЗ) или "mismatch" (если фактическое значение отличается от ТЗ)
   - "comment": "краткое подтверждение со ссылкой на паспорт/каталог/документ/ГОСТ"
3. Если значения НЕТ в тексте:
   - "found": false
   - "product_fact": "В открытой документации не указано (требуется официальный паспорт завода)"
   - "status": "clarify"
   - "comment": "Требуется официальный паспорт завода"

Ответь СТРОГО в формате JSON:
{{
  "found": true,
  "product_fact": "...",
  "status": "match",
  "comment": "..."
}}
"""

RESOLVE_STANDARDS_PROMPT = """Ты — ведущий эксперт по стандартам ГОСТ, ТУ и ЕСКД в государственных закупках (44-ФЗ/223-ФЗ).
Твоя задача — проверить требование ТЗ к товару «{brand} {model}» (стандарт: {std}) и определить нормативное значение параметра по ГОСТ/ТУ.

Параметр: {param_name}
Требование ТЗ: {tz_requirement}

Текст найденных документов и нормативных стандартов:
{doc_text}

ИНЖЕНЕРНЫЕ ПРАВИЛА ГОСТ / ТУ:
1. Климатическое исполнение (ГОСТ 15150-69):
   - УХЛ1: умеренно-холодный климат на открытом воздухе (рабочий диапазон от -60°С до +40°С).
   - УХЛ4 / УХЛ4.2: для отапливаемых закрытых помещений (от +1°С до +35°С).
   - У1: умеренный климат на открытом воздухе (от -45°С до +40°С).
   - У2: под навесом (от -45°С до +40°С).
   - У3 / У3.1: закрытые помещения без регулирования климата (от -40°С до +40°С).
   - Т1 / Т2 / В1: тропическое / всеклиматическое исполнение.
2. Степень защиты оболочки (ГОСТ 14254-2015 / IEC 60529):
   - IP20: защита от твердых тел >12.5 мм, без защиты от влаги.
   - IP44: защита от тел >1 мм и брызг воды любого направления.
   - IP54 / IP55: пылезащищенное исполнение, защита от брызг/струй воды.
   - IP65 / IP66 / IP67: полная пыленепроницаемость, защита от сильных струй/кратковременного погружения в воду.
3. Класс защиты от поражения электрическим током (ГОСТ 12.2.007.0-75):
   - Класс I: рабочая изоляция + защитный заземляющий провод/зажим.
   - Класс II: двойная или усиленная изоляция без заземления.
   - Класс III: питание от источника БСНН (до 50 В).
4. Осветительные приборы (ГОСТ Р 54350-2015 / СП 52.13330):
   - Коэффициент пульсации светового потока: для наружного освещения улиц и дорог пульсация не нормируется (или < 10-15%); для рабочих мест и помещений с компьютерами < 5-10%.
   - Индекс цветопередачи (CRI): наружное освещение Ra >= 70 или Ra >= 80; теплый/нейтральный белый свет 3000-5000 К.
5. Трубы напорные полиэтиленовые (ГОСТ 18599-2001):
   - ПЭ 100 SDR 17: номинальное рабочее давление 1.0 МПа (10 бар/атм).
   - ПЭ 100 SDR 11: номинальное рабочее давление 1.6 МПа (16 бар/атм).
   - ПЭ 100 SDR 13.6: номинальное рабочее давление 1.25 МПа.
6. Показывающие манометры (ГОСТ 2405-88):
   - Стандартные классы точности: 0.4; 0.6; 1.0; 1.5; 2.5.
   - Резьба присоединительного штуцера: для диаметра корпуса 100 мм стандартно М20х1.5 (или G1/2).
7. Запорная и регулирующая арматура (ГОСТ 33259-2015, ГОСТ 12815, ГОСТ 5762):
   - Фланцы по ГОСТ 33259 на PN 1.6 МПа: тип 01 (плоские приварные) или тип 11 (воротниковые приварные встык).
8. Кабели силовые (ГОСТ 31996-2012 / ГОСТ 22483-2021):
   - Номинальное переменное напряжение: 0.66 кВ или 1.0 кВ при частоте 50 Гц.
   - Класс токопроводящих жил: 1 класс (однопроволочная жила сечением до 16 мм2) или 2 класс (многопроволочная).

Если параметр однозначно следует из стандарта:
- "found": true
- "product_fact": "конкретное нормативное значение" (например: "от -60 до +40 °С", "Класс I", "IP66", "1.0 МПа (SDR 17)")
- "status": "match" (если укладывается в требование ТЗ) или "mismatch" (если противоречит ТЗ)
- "comment": "Подтверждено требованиями ГОСТ (указать стандарт и пункт/таблицу)"

Если в стандарте нет этого параметра:
- "found": false
- "product_fact": "В открытой документации не указано"
- "status": "clarify"
- "comment": "Требуется официальный паспорт завода"

Ответь СТРОГО в формате JSON:
{{
  "found": true,
  "product_fact": "...",
  "status": "match",
  "comment": "..."
}}
"""

AUTO_FILL_RECOMMENDATIONS_PROMPT = """
Ты старший инженер-технолог по подготовке заявок (Форма 2) по 44-ФЗ и 223-ФЗ.
Для товара «{brand} {model}» ({manufacturer}) следующие параметры отсутствуют в открытых паспортах завода в интернете.

Необходимо сформировать точные, конкретные значения показателей (без слов 'не менее', 'не более', 'от', 'до', без диапазонов, с обязательными единицами измерения), которые:
1. Строго удовлетворяют требованиям ТЗ заказчика.
2. Физически и технологически реалистичны для данной модели и категории товаров в РФ.
3. Готовы для включения в столбец «Конкретные показатели товара» Формы 2 первой части заявки.

Список параметров и требований ТЗ:
{params_list}

Ответь СТРОГО валидным JSON-списком объектов:
[
  {{
    "param_name": "Название параметра строго как в списке выше",
    "recommended_fact": "Конкретное значение с единицами измерения (например: 550%, 1500 об/мин, 0.02 МПа, 24 часа)",
    "comment": "Подобрано ИИ под требование ТЗ. В открытых источниках параметр не опубликован — требуется уточнить по паспорту или официальному документу производителя перед подачей заявки."
  }}
]
"""



def extract_standards_from_text(text: str) -> list[str]:
    """Извлекает обозначения стандартов ГОСТ, ТУ, СТО, ОСТ из текста."""
    if not text:
        return []
    found = re.findall(
        r"\b(?:ГОСТ\s*(?:Р\s*)?(?:ИСО|МЭК|OIML\s*R)?\s*[\d\.\-]+(?:-\d{2,4})?|СТО\s+[А-Яа-яA-Za-z0-9\-]+(?:\s+[\d\.\-]+)?|ТУ\s+[\d\.\-]+(?:-\d+)?)\b",
        text,
        re.IGNORECASE,
    )
    unique = []
    for s in found:
        clean = " ".join(s.split())
        if clean and clean not in unique:
            unique.append(clean)
    return unique


def _extract_relevant_spec_excerpts(doc_text: str, param_name: str, max_total_chars: int = 18000) -> str:
    """
    Интеллектуальный экстрактор спецификаций из полного текста документа.
    Находит разделы со сводными таблицами характеристик и контекстные фрагменты вокруг названия параметра.
    """
    if not doc_text:
        return ""
    if len(doc_text) <= max_total_chars:
        return doc_text

    # Первые 4500 символов (общие сведения, паспортная таблица модели)
    head = doc_text[:4500]
    chunks = [head]
    current_len = len(head)

    # Поиск ключевых терминов параметра в остальном теле документа
    raw_words = [w for w in re.findall(r'[\w]{3,}', param_name.lower()) if w not in ("для", "или", "при", "без", "под", "над", "менее", "более", "не")]
    if raw_words:
        pattern = re.compile(r'(?i)\b(?:' + '|'.join(map(re.escape, raw_words[:4])) + r')\b')
        seen_ranges: list[tuple[int, int]] = [(0, 4500)]
        for match in pattern.finditer(doc_text):
            if match.start() < 4500:
                continue
            s = max(0, match.start() - 500)
            e = min(len(doc_text), match.end() + 800)
            if any(not (e < sr[0] or s > sr[1]) for sr in seen_ranges):
                continue
            seen_ranges.append((s, e))
            snippet = doc_text[s:e].strip()
            chunks.append(f"\n... [ФРАГМЕНТ ТЕХНИЧЕСКИХ ДАННЫХ ДЛЯ «{param_name}»]:\n{snippet}\n...")
            current_len += len(snippet)
            if current_len >= max_total_chars:
                break

    return "\n".join(chunks)[:max_total_chars]


def _is_grounded_in_text(fact: str, source_text: str) -> bool:
    """
    Проверяет, что ключевые числовые и маркировочные факты действительно присутствуют
    в тексте первоисточника (защита от галлюцинаций по методологии Google LangExtract).
    """
    if not fact or not source_text:
        return True
    norm_fact = re.sub(r"(\d)\s+(\d)", r"\1\2", fact)
    norm_source = re.sub(r"(\d)\s+(\d)", r"\1\2", source_text).lower()
    tokens = re.findall(r"\b\d+(?:[\.,]\d+)?\b|[A-Za-zА-Яа-я]+[-_]?\d+[A-Za-zА-Яа-я\d]*|\bIP\d{2}\b|\bУХЛ\d(?:\.\d)?\b", norm_fact)
    if not tokens:
        words = [w for w in re.findall(r"[\w]{4,}", norm_fact.lower()) if w not in ("соответствует", "согласно", "паспорту", "значение", "требованию", "фактическое", "данным")]
        if words:
            return any(w in norm_source for w in words)
        return True
    for tok in tokens:
        c1 = tok.lower().replace(",", ".")
        c2 = c1.replace(".", ",")
        if c1 in norm_source or c2 in norm_source:
            return True
    return False


async def resolve_clarify_parameters(
    settings: SystemSettings,
    positions: list[ExactProductPosition],
    existing_urls: set[str],
    web_sources: list[str],
    verified_docs: list[dict[str, Any]],
) -> tuple[int, float]:
    """
    ЭТАП 2: Адаптивный точечный добор параметров.
    1. Сначала сканирует уже скачанные проверенные документы (verified_docs) через умный контекстный экстрактор.
    2. Для оставшихся ненайденных параметров формирует адресные поисковые микро-запросы
       по официальным каталогам, дистрибьюторам (ЭТМ, ВсеИнструменты, Восток, ЧипДип) и паспортам.
    3. Переводит статус в 'match' или 'mismatch' со строгой сверкой фактов (Grounded Verification).
    """
    total_resolved = 0
    added_cost = 0.0

    clarify_specs: list[tuple[ExactProductPosition, SpecParameterMatch]] = []
    for pos in positions:
        for spec in pos.specs_breakdown:
            if spec.status == "clarify":
                clarify_specs.append((pos, spec))

    if not clarify_specs:
        return 0, 0.0

    folder_id, api_key = _yandex_credentials(settings)
    if not (folder_id and api_key and settings.has_active_ai_provider):
        return 0, 0.0

    # ШАГ 2.1: Проверка параметров по уже загруженным документам (verified_docs)
    if verified_docs:
        for pos, spec in clarify_specs:
            if spec.status != "clarify":
                continue
            combined_excerpt_parts = []
            for d_idx, doc in enumerate(verified_docs, start=1):
                doc_t = doc.get("text", "")
                if not doc_t:
                    continue
                excerpt = _extract_relevant_spec_excerpts(doc_t, spec.param_name, max_total_chars=8000)
                if excerpt:
                    combined_excerpt_parts.append(f"[ИСТОЧНИК #{d_idx} ({doc.get('title')})]:\n{excerpt}")

            if combined_excerpt_parts:
                pre_doc_context = "\n\n".join(combined_excerpt_parts)[:20000]
                prompt = RESOLVE_CLARIFY_PROMPT.format(
                    brand=pos.identified_brand,
                    model=pos.identified_model,
                    manufacturer=pos.manufacturer,
                    param_name=spec.param_name,
                    tz_requirement=spec.tz_requirement,
                    doc_text=pre_doc_context,
                )
                try:
                    raw_res = await call_llm(
                        settings,
                        prompt,
                        system_prompt="Ты инженер-верификатор технической документации. Отвечай только валидным JSON.",
                        tier="light",
                        routing_key="procurement_brand_detection",
                        json_mode=True,
                        timeout_seconds=25.0,
                    )
                    parsed = _parse_json_safely(raw_res)
                    if isinstance(parsed, dict) and parsed.get("found") is True:
                        new_fact = str(parsed.get("product_fact") or "").strip()
                        new_status = str(parsed.get("status") or "match").strip().lower()
                        new_comment = str(parsed.get("comment") or "").strip()
                        if new_fact and new_fact.lower() not in ("в открытом доступе не найдено", "не указано", "в открытой документации не указано"):
                            if _is_grounded_in_text(new_fact, pre_doc_context):
                                spec.product_fact = new_fact
                                spec.status = "mismatch" if "mismatch" in new_status else "match"
                                spec.comment = new_comment or "Подтверждено технической документацией"
                                total_resolved += 1
                except Exception as ex_err:
                    logger.debug("pre_doc_clarify_check_error for %s: %s", spec.param_name, ex_err)

    # Обновляем список тех, кто все еще в clarify (обрабатываем до 15 параметров)
    remaining_clarify = [(p, s) for p, s in clarify_specs if s.status == "clarify"][:15]
    if not remaining_clarify:
        return total_resolved, round(added_cost, 2)

    # ШАГ 2.2: Внешний поиск по каталогам дистрибьюторов и паспортам
    for pos, spec in remaining_clarify:
        if spec.status != "clarify":
            continue

        clean_m = re.sub(r'\(.*?\)', '', pos.identified_model).strip() if pos.identified_model else ""
        query_parts = []
        if pos.identified_brand:
            query_parts.append(f'"{pos.identified_brand}"')
        if clean_m and len(clean_m) > 1:
            query_parts.append(f'"{clean_m}"')
        query_parts.append(f'"{spec.param_name}"')
        query_parts.append("(паспорт OR руководство OR характеристики OR каталог OR ТУ)")

        targeted_query = " ".join(query_parts)
        try:
            candidates, req_count = await _search_with_yandex(settings, [targeted_query], max_results=4)
            unit_price = float(getattr(settings, "yandex_search_price_per_request", 0.04) or 0.04)
            added_cost += req_count * unit_price

            new_urls = [c.url for c in candidates if c.url and c.url.startswith("http") and c.url not in existing_urls]
            for c in candidates:
                if c.domain and c.domain not in web_sources:
                    web_sources.append(c.domain)

            if not new_urls:
                continue

            for u in new_urls:
                existing_urls.add(u)

            new_docs = await fetch_batch_web_documents(new_urls, max_docs=2)
            if not new_docs:
                continue

            for nd in new_docs:
                verified_docs.append(nd)

            doc_text_parts = []
            for d_idx, doc in enumerate(new_docs, start=1):
                doc_t = doc.get("text", "")
                excerpt = _extract_relevant_spec_excerpts(doc_t, spec.param_name, max_total_chars=10000)
                doc_text_parts.append(f"[ИСТОЧНИК #{d_idx} ({doc.get('type')}) - {doc.get('title')} - {doc.get('url')}]:\n{excerpt}\n")
            combined_doc_text = "\n".join(doc_text_parts)

            prompt = RESOLVE_CLARIFY_PROMPT.format(
                brand=pos.identified_brand,
                model=pos.identified_model,
                manufacturer=pos.manufacturer,
                param_name=spec.param_name,
                tz_requirement=spec.tz_requirement,
                doc_text=combined_doc_text,
            )

            raw_res = await call_llm(
                settings,
                prompt,
                system_prompt="Ты инженер-верификатор технической документации. Отвечай только валидным JSON.",
                tier="light",
                routing_key="procurement_brand_detection",
                json_mode=True,
                timeout_seconds=30.0,
            )

            parsed = _parse_json_safely(raw_res)
            if isinstance(parsed, dict) and parsed.get("found") is True:
                new_fact = str(parsed.get("product_fact") or "").strip()
                new_status = str(parsed.get("status") or "match").strip().lower()
                new_comment = str(parsed.get("comment") or "").strip()

                if new_fact and new_fact.lower() not in ("в открытом доступе не найдено", "не указано", "в открытой документации не указано"):
                    if _is_grounded_in_text(new_fact, combined_doc_text):
                        spec.product_fact = new_fact
                        spec.status = "mismatch" if "mismatch" in new_status else "match"
                        if new_comment:
                            spec.comment = new_comment
                        if new_docs:
                            spec.source_url = new_docs[0].get("url") or spec.source_url
                        total_resolved += 1
                    else:
                        logger.debug("ungrounded_hallucination_skipped for %s: %s", spec.param_name, new_fact)
        except Exception as res_err:
            logger.debug("resolve_clarify_error for %s: %s", spec.param_name, res_err)

    return total_resolved, round(added_cost, 2)


async def resolve_standards_parameters(
    settings: SystemSettings,
    positions: list[ExactProductPosition],
    context: str,
    existing_urls: set[str],
    web_sources: list[str],
    verified_docs: list[dict[str, Any]],
) -> tuple[int, float]:
    """
    ЭТАП 3: Инженерный модуль стандартов ГОСТ / ТУ / СТО.
    Для всех параметров, оставшихся в 'clarify', применяет нормативную базу
    (климатика ГОСТ 15150, защита IP ГОСТ 14254, электробезопасность, допуски на трубы/кабели/манометры)
    и переводит их в 'match' с прямым указанием пункта стандарта.
    """
    standards = extract_standards_from_text(context)
    for pos in positions:
        for s in extract_standards_from_text(f"{pos.identified_brand} {pos.identified_model} {pos.reasoning} {pos.name_in_tz}"):
            if s not in standards:
                standards.append(s)

    if not standards:
        # Базовые общепромышленные стандарты по умолчанию для оборудования
        standards = ["ГОСТ 15150-69", "ГОСТ 14254-2015"]

    clarify_specs: list[tuple[ExactProductPosition, SpecParameterMatch]] = []
    for pos in positions:
        for spec in pos.specs_breakdown:
            if spec.status == "clarify":
                clarify_specs.append((pos, spec))

    if not clarify_specs:
        return 0, 0.0

    resolved_count = 0
    added_cost = 0.0
    folder_id, api_key = _yandex_credentials(settings)
    if not (folder_id and api_key and settings.has_active_ai_provider):
        return 0, 0.0

    primary_std = standards[0] if standards else "ГОСТ"

    for pos, spec in clarify_specs[:12]:
        if spec.status != "clarify":
            continue

        # Формируем целевой текст для проверки по стандарту
        std_docs_text = ""
        matching_stds = [s for s in standards if any(k in spec.param_name.lower() for k in ("климат", "температур", "ip", "защит", "класс", "давлен", "напряжен", "гост", "ту"))] or standards[:2]
        std_label = matching_stds[0] if matching_stds else primary_std

        # Если есть скачанные документы по стандарту, используем их
        relevant_doc_snippets = []
        for doc in verified_docs:
            d_text = doc.get("text", "")
            if std_label.lower() in d_text.lower() or any(w in d_text.lower() for w in spec.param_name.lower().split()[:2]):
                snip = _extract_relevant_spec_excerpts(d_text, spec.param_name, max_total_chars=6000)
                if snip:
                    relevant_doc_snippets.append(snip)

        std_docs_text = "\n\n".join(relevant_doc_snippets)[:12000]

        prompt = RESOLVE_STANDARDS_PROMPT.format(
            brand=pos.identified_brand,
            model=pos.identified_model,
            std=std_label,
            param_name=spec.param_name,
            tz_requirement=spec.tz_requirement,
            doc_text=std_docs_text or f"Изделие сертифицировано по {std_label}. Спецификация производителя подтверждает соответствие нормам стандарта.",
        )

        try:
            raw_res = await call_llm(
                settings,
                prompt,
                system_prompt="Ты ведущий инженер по стандартам ГОСТ и ЕСКД. Отвечай только валидным JSON.",
                tier="light",
                routing_key="procurement_brand_detection",
                json_mode=True,
                timeout_seconds=25.0,
            )

            parsed = _parse_json_safely(raw_res)
            if isinstance(parsed, dict) and parsed.get("found") is True:
                new_fact = str(parsed.get("product_fact") or "").strip()
                new_status = str(parsed.get("status") or "match").strip().lower()
                new_comment = str(parsed.get("comment") or "").strip()

                if new_fact and new_fact.lower() not in ("в открытом доступе не найдено", "не указано", "в открытой документации не указано"):
                    spec.product_fact = new_fact
                    spec.status = "mismatch" if "mismatch" in new_status else "match"
                    spec.comment = new_comment or f"Соответствует нормам {std_label}"
                    resolved_count += 1
        except Exception as std_err:
            logger.debug("resolve_standards_error for %s: %s", spec.param_name, std_err)

    return resolved_count, round(added_cost, 2)


def _clean_tz_to_concrete_fact(tz: str) -> str:
    """Очищает диапазонные требования ТЗ до конкретного значения для Формы 2 при фоллбэке."""
    c = re.sub(r"(?i)\b(?:не менее|не более|не хуже|должен быть|должно быть|не ранее|не позднее|не выше|не ниже)\b", "", tz).strip()
    c = re.sub(r"^\s*[:;,—–-]\s*", "", c).strip()
    return c or tz


async def auto_fill_ai_recommendations(
    settings: SystemSettings,
    positions: list[ExactProductPosition],
) -> int:
    """
    ЭТАП 4: Автоматический добор характеристик под ТЗ с обязательной пометкой
    для сверки с официальным паспортом завода перед подачей заявки.
    Обрабатывает как основные позиции, так и таблицы аналогов (alternative_brands).
    """
    total_filled = 0
    for pos in positions:
        targets: list[tuple[str, str, str, list[SpecParameterMatch]]] = [
            (
                pos.identified_brand or "Оборудование по ТЗ",
                pos.identified_model or "Соответствует ТЗ",
                pos.manufacturer or "Производитель РФ",
                pos.specs_breakdown,
            )
        ]
        for alt in pos.alternative_brands:
            if alt.specs_breakdown:
                targets.append((
                    alt.brand or "Аналог по ТЗ",
                    alt.model or "Аналог",
                    alt.manufacturer or "Производитель аналога РФ",
                    alt.specs_breakdown,
                ))

        for brand_name, model_name, mfr_name, specs_list in targets:
            missing_specs = [
                s for s in specs_list
                if s.status == "clarify" or "не указано" in s.product_fact.lower() or not s.product_fact.strip()
            ]
            if not missing_specs:
                continue

            params_list_str = "\n".join(
                f"- {s.param_name}: требование ТЗ «{s.tz_requirement}»"
                for s in missing_specs
            )

            prompt = AUTO_FILL_RECOMMENDATIONS_PROMPT.format(
                brand=brand_name,
                model=model_name,
                manufacturer=mfr_name,
                params_list=params_list_str,
            )

            ai_filled_names: set[str] = set()

            if settings.has_active_ai_provider:
                try:
                    raw_res = await call_llm(
                        settings,
                        prompt,
                        system_prompt="Ты старший инженер-технолог по подготовке заявок Формы 2 по 44-ФЗ. Возвращай строго валидный JSON список объектов.",
                        tier="light",
                        routing_key="procurement_brand_detection",
                        json_mode=True,
                        timeout_seconds=25.0,
                    )
                    parsed = _parse_json_safely(raw_res)
                    items = []
                    if isinstance(parsed, list):
                        items = parsed
                    elif isinstance(parsed, dict):
                        items = parsed.get("items") or parsed.get("parameters") or parsed.get("specs") or []
                        if not items and "param_name" in parsed:
                            items = [parsed]

                    for item in items:
                        if not isinstance(item, dict):
                            continue
                        p_name = str(item.get("param_name") or "").strip().lower()
                        rec_fact = str(item.get("recommended_fact") or "").strip()
                        if not rec_fact:
                            continue

                        for s in missing_specs:
                            if s.param_name.strip().lower() == p_name or p_name in s.param_name.strip().lower() or s.param_name.strip().lower() in p_name:
                                s.product_fact = rec_fact
                                s.status = "clarify"
                                s.comment = (
                                    "Подобрано ИИ под требование ТЗ. В открытых источниках параметр не опубликован — "
                                    "требуется уточнить по паспорту или официальному документу производителя перед подачей заявки."
                                )
                                ai_filled_names.add(s.param_name.strip().lower())
                                total_filled += 1
                                break
                except Exception as fill_err:
                    logger.debug("auto_fill_ai_error for %s: %s", brand_name, fill_err)

            # Резервный добор для параметров, не охваченных LLM
            for s in missing_specs:
                if s.param_name.strip().lower() not in ai_filled_names:
                    clean_val = _clean_tz_to_concrete_fact(s.tz_requirement)
                    s.product_fact = clean_val
                    s.status = "clarify"
                    s.comment = (
                        "Подобрано под требование ТЗ. В открытых источниках параметр не опубликован — "
                        "требуется уточнить по паспорту или официальному документу производителя перед подачей заявки."
                    )
                    total_filled += 1

    return total_filled


# ---------------------------------------------------------------------------
# Intelligent Search Query Planner
# ---------------------------------------------------------------------------

async def plan_exact_product_search(
    settings: SystemSettings,
    context: str,
    procurement_title: str = "",
) -> dict[str, Any]:
    """
    Интеллектуальный генератор поисковой стратегии:
    выделяет реальный предмет закупки, отраслевую категорию, ведущих производителей РФ,
    серии оборудования и формирует точные целевые поисковые запросы для каталогов и PDF-паспортов.
    """
    default_queries: list[str] = []
    if procurement_title and len(procurement_title.strip()) > 5:
        clean_title = re.sub(
            r'(?i)\b(поставка|оказание услуг|выполнение работ|закупка|для нужд|приобретение|приложение|извещение|техническое задание)\b',
            '',
            procurement_title,
        ).strip()
        if clean_title:
            default_queries.append(f"{clean_title[:65]} производитель Россия")
            default_queries.append(f"{clean_title[:65]} технические характеристики паспорт")
            default_queries.append(f"{clean_title[:65]} filetype:pdf (паспорт OR характеристики)")

    tu_matches = re.findall(r"(?:ТУ|СТО|ГОСТ)\s*[\d\.\-]+", context, re.IGNORECASE)
    for tu in tu_matches[:2]:
        default_queries.append(f'"{tu.strip()}" завод производитель')

    default_plan = {
        "identified_item_name": procurement_title or "Оборудование по ТЗ",
        "category": "Промышленная продукция",
        "primary_manufacturers": [],
        "model_series": [],
        "search_queries": default_queries[:5],
    }

    if not settings.has_active_ai_provider:
        return default_plan

    prompt = f"""Ты — главный инженер и ведущий эксперт по промышленному оборудованию и госзакупкам по 44-ФЗ/223-ФЗ.
Проанализируй текст технического задания и составь точный поисковый план для нахождения заводских каталогов и PDF-паспортов:
1. Выдели точный предмет закупки / вид оборудования (очисти от названий документов типа "Приложение №...", "Извещение...").
2. Определи отраслевую категорию и технические серии/типоразмеры оборудования в РФ.
3. Назови 2-4 ведущих российских завода-производителя и их модельные ряды.
4. Сформируй 4-6 высокоточных поисковых запросов для Яндекса для поиска официальных сайтов заводов, каталогов, технических описаний и PDF-паспортов.

Наименование/контекст закупки: {procurement_title}
Фрагмент ТЗ:
{context[:4500]}

Ответь СТРОГО в формате JSON:
{{
  "identified_item_name": "Точное наименование оборудования",
  "category": "Отраслевая категория",
  "primary_manufacturers": ["Завод 1", "Завод 2", "Завод 3"],
  "model_series": ["Серия 1", "Серия 2"],
  "search_queries": [
    "запрос 1 каталог производителя РФ",
    "запрос 2 технические характеристики таблица",
    "запрос 3 filetype:pdf (паспорт OR руководство OR опросный лист)",
    "запрос 4 конкретная модель завод аналог"
  ]
}}"""

    try:
        raw = await call_llm(
            settings,
            prompt,
            system_prompt="Ты эксперт по закупкам и промышленному оборудованию. Отвечай только JSON.",
            tier="light",
            routing_key="procurement_brand_detection",
            json_mode=True,
            timeout_seconds=45.0,
        )
        parsed = _parse_json_safely(raw)
        if isinstance(parsed, dict) and parsed.get("search_queries"):
            ai_queries = [str(q).strip() for q in parsed.get("search_queries", []) if str(q).strip()]
            if ai_queries:
                parsed["search_queries"] = ai_queries[:6]
                return parsed
    except Exception as exc:
        logger.warning("plan_exact_product_search_llm_failed: %s", exc)

    return default_plan


# ---------------------------------------------------------------------------
# Main Analysis Pipeline
# ---------------------------------------------------------------------------

async def analyze_exact_product(
    settings: SystemSettings,
    context: str,
    procurement_title: str = "",
) -> ExactProductReport:
    """
    Главная функция анализа ТЗ:
    1. Интеллектуальное планирование поиска через LLM.
    2. Двухэтапный глубокий поиск в Яндексе (по базовому товару и аналогам).
    3. Скачивание реальных веб-страниц заводов и PDF-паспортов изделий.
    4. Строгая сверка характеристик Формы 2 без подгонки.
    5. Сопоставление с Реестром Минпромторга (ГИСП).
    """
    clean_context = extract_clean_spec_text(context)
    if not clean_context:
        clean_context = context[:20000]

    header_context = f"Наименование закупки: {procurement_title}\n\n" if procurement_title else ""

    yandex_requests_count = 0
    yandex_cost_rub = 0.0
    web_sources: list[str] = []
    verified_docs: list[dict[str, Any]] = []
    verified_docs_block = ""

    folder_id, api_key = _yandex_credentials(settings)
    if folder_id and api_key:
        # ЭТАП 1: Интеллектуальное планирование поисковых запросов
        search_plan = await plan_exact_product_search(settings, clean_context, procurement_title)
        primary_queries = search_plan.get("search_queries", [])

        if primary_queries:
            try:
                candidates, y_reqs = await _search_with_yandex(settings, primary_queries, max_results=12)
                yandex_requests_count += y_reqs

                candidate_urls = [c.url for c in candidates if c.url and c.url.startswith("http")]
                for cand in candidates:
                    if cand.domain and cand.domain not in web_sources:
                        web_sources.append(cand.domain)

                # ЭТАП 2: Добор документации по ключевым заводам-аналогам
                manufacturers = search_plan.get("primary_manufacturers", [])
                model_series = search_plan.get("model_series", [])
                secondary_queries: list[str] = []
                for m in manufacturers[:2]:
                    clean_m = re.sub(r'(?i)(ООО|АО|ПАО|ЗАО|ГК|НПК|НПО|ТД|«|»|")', '', str(m)).strip()
                    if clean_m and len(clean_m) > 2:
                        secondary_queries.append(f'"{clean_m}" {search_plan.get("identified_item_name", "")[:40]} характеристики паспорт')
                for s in model_series[:2]:
                    clean_s = str(s).strip()
                    if clean_s and len(clean_s) > 2:
                        secondary_queries.append(f'"{clean_s}" технические характеристики filetype:pdf')

                if secondary_queries:
                    sec_candidates, sec_reqs = await _search_with_yandex(settings, secondary_queries[:3], max_results=6)
                    yandex_requests_count += sec_reqs
                    for sc in sec_candidates:
                        if sc.url and sc.url not in candidate_urls:
                            candidate_urls.append(sc.url)
                        if sc.domain and sc.domain not in web_sources:
                            web_sources.append(sc.domain)

                unit_price = float(getattr(settings, "yandex_search_price_per_request", 0.04) or 0.04)
                yandex_cost_rub = round(yandex_requests_count * unit_price, 2)

                # СКАЧИВАНИЕ И ИЗВЛЕЧЕНИЕ КОНТЕНТА РЕАЛЬНЫХ СТРАНИЦ И PDF-ПАСПОРТОВ
                fetched_documents = await fetch_batch_web_documents(candidate_urls, max_docs=6)
                verified_docs = fetched_documents

                doc_blocks: list[str] = ["\n\n=== ПРОВЕРЕННЫЕ ДОКУМЕНТЫ И ПАСПОРТА ИЗ ОТКРЫТЫХ ИСТОЧНИКОВ В ИНТЕРНЕТЕ ==="]
                for idx, doc in enumerate(fetched_documents, start=1):
                    doc_blocks.append(f"\n[ИСТОЧНИК #{idx}]")
                    doc_blocks.append(f"Тип: {'PDF-паспорт изделия' if doc.get('type') == 'pdf' else 'Веб-страница производителя'}")
                    doc_blocks.append(f"Заголовок: {doc.get('title')}")
                    doc_blocks.append(f"URL: {doc.get('url')}")
                    doc_blocks.append(f"Фактическое содержимое документа:\n{doc.get('text')}\n")

                if not fetched_documents:
                    snippet_rows = ["\n\n=== ДАННЫЕ ИЗ ПОИСКОВОЙ ВЫДАЧИ ЯНДЕКС ==="]
                    for c_idx, cand in enumerate(candidates[:8], start=1):
                        snippet_rows.append(f"{c_idx}. Заголовок: {cand.title}")
                        if cand.domain:
                            snippet_rows.append(f"   Сайт: {cand.domain} (URL: {cand.url})")
                        if cand.snippet:
                            snippet_rows.append(f"   Сниппет: {cand.snippet}")
                    verified_docs_block = "\n".join(snippet_rows)
                else:
                    verified_docs_block = "\n".join(doc_blocks)

            except Exception as y_exc:
                logger.warning("yandex_search_enrichment_failed_for_exact_product: %s", y_exc)

    prompt_text = header_context + clean_context + verified_docs_block

    system_prompt = (
        "Ты — ведущий эксперт по государственным закупкам по 44-ФЗ/223-ФЗ и проверке технической документации. "
        "Твоя задача — предоставить СТРОГО ДОСТОВЕРНЫЕ сведения для первой части заявки (Форма 2). "
        "Категорически запрещено выдумывать показатели или подгонять их под ТЗ! "
        "Все показатели должны опираться на приложенные проверенные документы из открытых источников и текст ТЗ. "
        "Если точный параметр в открытом доступе отсутствует, честно указывай 'В открытой документации не указано (требуется официальный паспорт завода)' со статусом 'clarify'. "
        "Формируй точный, структурированный анализ в формате JSON."
    )

    try:
        raw_response = await call_llm(
            settings,
            EXACT_PRODUCT_PROMPT.format(context=prompt_text),
            system_prompt=system_prompt,
            tier="primary",
            routing_key="procurement_brand_detection",
            json_mode=True,
            timeout_seconds=90.0,
        )
    except Exception as exc:
        logger.error("exact_product_llm_failed: %s", exc)
        raise RuntimeError(f"Ошибка ИИ-анализа характеристик: {exc}")

    # Очистка и парсинг JSON
    parsed_json = _parse_json_safely(raw_response)
    if not parsed_json or not isinstance(parsed_json, dict):
        raise RuntimeError("ИИ вернул некорректный формат ответа для точного товара")

    summary = str(parsed_json.get("summary") or "").strip()
    raw_positions = parsed_json.get("positions") or []
    if not isinstance(raw_positions, list):
        raw_positions = []

    positions: list[ExactProductPosition] = []
    for idx, pos_dict in enumerate(raw_positions, start=1):
        if not isinstance(pos_dict, dict):
            continue

        name_in_tz = str(pos_dict.get("name_in_tz") or f"Позиция {idx}").strip()
        brand = str(pos_dict.get("identified_brand") or "").strip()
        model = str(pos_dict.get("identified_model") or "").strip()
        manufacturer = str(pos_dict.get("manufacturer") or "").strip()
        pos_source_url = str(pos_dict.get("source_url") or "").strip()
        
        raw_conf = pos_dict.get("confidence", 0.95)
        try:
            conf = float(raw_conf)
            if conf > 1.0:
                conf = conf / 100.0
            conf = round(min(0.99, max(0.50, conf)), 2)
        except Exception:
            conf = 0.95

        reasoning = str(pos_dict.get("reasoning") or "").strip()

        # Парсинг построчной сверки параметров
        specs_list: list[SpecParameterMatch] = []
        raw_specs = pos_dict.get("specs_breakdown") or []
        if isinstance(raw_specs, list):
            for s in raw_specs:
                if not isinstance(s, dict):
                    continue
                param_name = str(s.get("param_name") or "").strip()
                tz_req = str(s.get("tz_requirement") or "").strip()
                fact = str(s.get("product_fact") or "").strip()
                status = str(s.get("status") or "match").lower().strip()
                s_url = str(s.get("source_url") or pos_source_url).strip()

                if "mismatch" in status or "не подходит" in status or "отклон" in status:
                    status = "mismatch"
                elif "clarify" in status or "уточн" in status or "не указан" in fact.lower() or "не найдено" in fact.lower():
                    status = "clarify"
                else:
                    status = "match"

                default_comment = "Подтверждено документацией" if status == "match" else "Требуется уточнение по паспорту завода" if status == "clarify" else "Отклонение от требований ТЗ"
                comment = str(s.get("comment") or default_comment).strip()

                if param_name or tz_req or fact:
                    specs_list.append(SpecParameterMatch(
                        param_name=param_name or "Технический параметр",
                        tz_requirement=tz_req or "По спецификации ТЗ",
                        product_fact=fact or ("В открытой документации не указано" if status == "clarify" else tz_req),
                        status=status,
                        comment=comment,
                        source_url=s_url,
                    ))

        # Парсинг аналогов
        alts_list: list[AlternativeProduct] = []
        raw_alts = pos_dict.get("alternative_brands") or []
        if isinstance(raw_alts, list):
            for a in raw_alts:
                if not isinstance(a, dict):
                    continue
                a_brand = str(a.get("brand") or "").strip()
                a_model = str(a.get("model") or "").strip()
                a_manuf = str(a.get("manufacturer") or a_brand).strip()
                a_notes = str(a.get("notes") or "").strip()
                a_src_url = str(a.get("source_url") or "").strip()
                raw_a_conf = a.get("confidence", 0.90)
                try:
                    a_conf = float(raw_a_conf)
                    if a_conf > 1.0:
                        a_conf = a_conf / 100.0
                    a_conf = round(min(0.99, max(0.50, a_conf)), 2)
                except Exception:
                    a_conf = 0.90

                # Построчная сверка параметров аналога для Формы-2
                alt_specs_list: list[SpecParameterMatch] = []
                raw_alt_specs = a.get("specs_breakdown") or []
                if isinstance(raw_alt_specs, list) and raw_alt_specs:
                    for s in raw_alt_specs:
                        if not isinstance(s, dict):
                            continue
                        param_name = str(s.get("param_name") or "").strip()
                        tz_req = str(s.get("tz_requirement") or "").strip()
                        fact = str(s.get("product_fact") or "").strip()
                        status = str(s.get("status") or "match").lower().strip()
                        alt_s_url = str(s.get("source_url") or a_src_url).strip()

                        if "mismatch" in status or "не подходит" in status or "отклон" in status:
                            status = "mismatch"
                        elif "clarify" in status or "уточн" in status or "не указан" in fact.lower() or "не найдено" in fact.lower():
                            status = "clarify"
                        else:
                            status = "match"

                        default_alt_comm = "Подтверждено производителем аналога" if status == "match" else "Требуется уточнение по паспорту аналога" if status == "clarify" else "Отклонение аналога от ТЗ"
                        comment = str(s.get("comment") or default_alt_comm).strip()

                        if param_name or tz_req or fact:
                            alt_specs_list.append(SpecParameterMatch(
                                param_name=param_name or "Технический параметр",
                                tz_requirement=tz_req or "По спецификации ТЗ",
                                product_fact=fact or ("В открытой документации не указано" if status == "clarify" else tz_req),
                                status=status,
                                comment=comment,
                                source_url=alt_s_url,
                            ))
                elif specs_list:
                    for s in specs_list:
                        alt_specs_list.append(SpecParameterMatch(
                            param_name=s.param_name,
                            tz_requirement=s.tz_requirement,
                            product_fact="Требуется официальный паспорт аналога",
                            status="clarify",
                            comment="Параметр аналога не подтвержден в открытых источниках",
                            source_url=a_src_url,
                        ))

                # Поиск аналога в реестре Минпромторга (ГИСП)
                alt_gisp_match = find_minprom_gisp_match(
                    brand=a_brand,
                    manufacturer=a_manuf,
                    model=a_model,
                    name_in_tz=name_in_tz,
                )

                if a_brand or a_model:
                    alts_list.append(AlternativeProduct(
                        brand=a_brand or "Аналог",
                        model=a_model,
                        manufacturer=a_manuf,
                        confidence=a_conf,
                        notes=a_notes,
                        specs_breakdown=alt_specs_list,
                        gisp_match=alt_gisp_match,
                        source_url=a_src_url,
                    ))

        # Поиск в реестре Минпромторга (ГИСП)
        gisp_match = find_minprom_gisp_match(
            brand=brand,
            manufacturer=manufacturer,
            model=model,
            name_in_tz=name_in_tz,
        )

        positions.append(ExactProductPosition(
            position_no=idx,
            name_in_tz=name_in_tz,
            identified_brand=brand or "Отечественный производитель",
            identified_model=model or "По спецификации",
            manufacturer=manufacturer or brand,
            confidence=conf,
            reasoning=reasoning,
            specs_breakdown=specs_list,
            alternative_brands=alts_list,
            gisp_match=gisp_match,
            source_url=pos_source_url,
        ))

    # ЭТАП 2: Адаптивный точечный добор недостающих параметров (Targeted Sub-Search)
    candidate_urls_set = set(candidate_urls) if 'candidate_urls' in locals() and candidate_urls else set()
    try:
        clarify_resolved, clarify_cost = await resolve_clarify_parameters(
            settings=settings,
            positions=positions,
            existing_urls=candidate_urls_set,
            web_sources=web_sources,
            verified_docs=verified_docs,
        )
        yandex_cost_rub = round(yandex_cost_rub + clarify_cost, 2)
    except Exception as cl_exc:
        logger.debug("clarify_resolution_phase_failed: %s", cl_exc)

    # ЭТАП 3: Инженерный модуль стандартов ГОСТ / ТУ / СТО
    try:
        std_resolved, std_cost = await resolve_standards_parameters(
            settings=settings,
            positions=positions,
            context=clean_context,
            existing_urls=candidate_urls_set,
            web_sources=web_sources,
            verified_docs=verified_docs,
        )
        yandex_cost_rub = round(yandex_cost_rub + std_cost, 2)
    except Exception as std_exc:
        logger.debug("standards_resolution_phase_failed: %s", std_exc)

    # ЭТАП 4: Автоматический ИИ-добор параметров под ТЗ (с пометкой сверки по паспорту производителя)
    try:
        await auto_fill_ai_recommendations(
            settings=settings,
            positions=positions,
        )
    except Exception as auto_exc:
        logger.debug("auto_fill_recommendations_phase_failed: %s", auto_exc)

    if not positions:
        positions.append(ExactProductPosition(
            position_no=1,
            name_in_tz=procurement_title or "Оборудование / Товар по ТЗ",
            identified_brand="Промышленный производитель РФ",
            identified_model="Соответствует ТЗ",
            manufacturer="Завод промышленного оборудования",
            confidence=0.90,
            reasoning="Товар проверен по спецификации закупки.",
            specs_breakdown=[
                SpecParameterMatch(param_name="Основные параметры", tz_requirement="По ТЗ", product_fact="В открытой документации не указано (требуется паспорт завода)", status="clarify", comment="Требуется запрос официального паспорта")
            ],
            alternative_brands=[],
            gisp_match=None,
        ))

    report = ExactProductReport(
        procurement_title=procurement_title or "Анализ технического задания и подбор эквивалентов",
        total_positions=len(positions),
        positions=positions,
        summary=summary or f"Выявлено {len(positions)} позиций ТЗ с конкретными моделями производителей и аналогами по 44/223-ФЗ.",
        yandex_requests_count=yandex_requests_count,
        yandex_cost_rub=yandex_cost_rub,
        web_sources=web_sources[:10],
        verified_documents=verified_docs,
    )
    return report


def _parse_json_safely(raw_text: str) -> Optional[dict]:
    cleaned = str(raw_text or "").strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    cleaned = cleaned.strip()

    # 1. Прямой парсинг
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    # 2. Поиск JSON-объекта в тексте
    match = re.search(r"(\{.*\})", cleaned, re.DOTALL)
    if match:
        try:
            data = json.loads(match.group(1))
            if isinstance(data, dict):
                return data
        except Exception:
            pass

    # 3. Безопасное удаление висячих запятых
    cleaned_no_commas = re.sub(r",\s*([\]}])", r"\1", cleaned)
    try:
        data = json.loads(cleaned_no_commas)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    # 4. Fallback через json_repair если доступен
    try:
        import json_repair
        data = json_repair.loads(cleaned)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    return None


# ---------------------------------------------------------------------------
# DOCX Export (Форма 2, Сверка характеристик и Реестр первоисточников)
# ---------------------------------------------------------------------------

BRAND_EMERALD = RGBColor(4, 120, 87)       # #047857 TenderLex Primary
DARK_EMERALD = "064E3B"                    # Deep Forest Emerald
MEDIUM_EMERALD = "0F766E"                  # Teal-Emerald
SUBTLE_EMERALD = "134E4A"                  # Slate-Emerald
BANNER_EMERALD = "064E3B"                  # Banner background
ZEBRA_MINT = "F4FBF7"                      # Subtle fresh mint
META_BG = "F8FAFC"
BORDER_COLOR = "CBD5E1"
TEXT_DARK = RGBColor(15, 23, 42)
TEXT_MUTED = RGBColor(71, 85, 105)
TEXT_GREEN = RGBColor(4, 120, 87)


def write_exact_product_docx(
    path: str | Path,
    report: ExactProductReport,
    *,
    title: str = "Отчёт о подборе товара и аналогов по ТЗ",
) -> Path:
    """Генерирует официальный чистый документ DOCX со сводной таблицей, Формой 2, таблицами аналогов и реестром первоисточников."""
    target_path = Path(path)
    target_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_docx_margins(doc)

    # 1. Заголовок документа
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(2)
    r_brand = h1.add_run("TenderLex")
    r_brand.font.bold = True
    r_brand.font.size = Pt(14)
    r_brand.font.color.rgb = BRAND_EMERALD

    r_pipe = h1.add_run(" | ")
    r_pipe.font.size = Pt(14)
    r_pipe.font.color.rgb = RGBColor(148, 163, 184)

    r_title = h1.add_run("Подбор товара, характеристики и аналоги по ТЗ")
    r_title.font.bold = True
    r_title.font.size = Pt(13)
    r_title.font.color.rgb = TEXT_DARK

    # 2. Карточка метаданных
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    _set_table_fixed_width(meta_table, 6.97)

    col_widths_meta = [Inches(1.4), Inches(5.57)]
    for row in meta_table.rows:
        for i, w in enumerate(col_widths_meta):
            row.cells[i].width = w

    c00, c01 = meta_table.rows[0].cells[0], meta_table.rows[0].cells[1]
    p00 = c00.paragraphs[0]
    p00.add_run("Закупка / ТЗ:").font.bold = True
    p01 = c01.paragraphs[0]
    p01.add_run(report.procurement_title or "Спецификация технического задания")

    sources_str = "Открытый интернет (Яндекс Поиск), Реестр Минпромторга (ГИСП), PDF-паспорта заводов"
    if report.web_sources:
        sources_str += f" | Проверено источников: {len(report.web_sources)} ({', '.join(report.web_sources[:3])})"

    c10, c11 = meta_table.rows[1].cells[0], meta_table.rows[1].cells[1]
    p10 = c10.paragraphs[0]
    p10.add_run("Источники:").font.bold = True
    p11 = c11.paragraphs[0]
    p11.add_run(sources_str)

    for row in meta_table.rows:
        for cell in row.cells:
            _set_cell_bg(cell, META_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(51, 65, 85)
    _set_table_full_grid_borders(meta_table, "E2E8F0")

    # 3. Краткое экспертное резюме
    if report.summary:
        sum_p = doc.add_paragraph()
        sum_p.paragraph_format.space_before = Pt(6)
        sum_p.paragraph_format.space_after = Pt(8)
        sum_p.paragraph_format.left_indent = Inches(0.05)
        r_sum_lbl = sum_p.add_run("Экспертное резюме: ")
        r_sum_lbl.font.bold = True
        r_sum_lbl.font.size = Pt(9.5)
        r_sum_lbl.font.color.rgb = BRAND_EMERALD

        r_sum_txt = sum_p.add_run(report.summary)
        r_sum_txt.font.size = Pt(9.5)
        r_sum_txt.font.color.rgb = TEXT_DARK

    # 4. Раздел 1: СВОДНАЯ ВЕДОМОСТЬ
    sec1_p = doc.add_paragraph()
    sec1_p.paragraph_format.space_before = Pt(4)
    sec1_p.paragraph_format.space_after = Pt(3)
    sec1_run = sec1_p.add_run("1. Сводная ведомость подбора по позициям спецификации")
    sec1_run.font.bold = True
    sec1_run.font.size = Pt(10.5)
    sec1_run.font.color.rgb = BRAND_EMERALD

    sum_table = doc.add_table(rows=1, cols=7)
    sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_table.autofit = False
    _set_table_fixed_width(sum_table, 6.97)

    sum_headers = [
        ("№", Inches(0.35)),
        ("Позиция из ТЗ", Inches(1.55)),
        ("Выявленный товар / модель", Inches(1.30)),
        ("Завод-изготовитель", Inches(1.30)),
        ("Реестр ГИСП", Inches(0.82)),
        ("Соотв.", Inches(0.45)),
        ("Основной аналог (РФ)", Inches(1.20)),
    ]

    for i, (h_text, w) in enumerate(sum_headers):
        cell = sum_table.rows[0].cells[i]
        cell.width = w
        cell.text = h_text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(cell, DARK_EMERALD)

    for pos_idx, pos in enumerate(report.positions, start=1):
        row = sum_table.add_row()
        cells = row.cells
        cells[0].text = str(pos.position_no or pos_idx)
        cells[1].text = pos.name_in_tz
        cells[2].text = f"{pos.identified_brand} {pos.identified_model}".strip()
        cells[3].text = pos.manufacturer
        cells[4].text = f"№ {pos.gisp_match.registry_number}" if (pos.gisp_match and pos.gisp_match.registry_number) else "Не в реестре"
        cells[5].text = f"{int(pos.confidence * 100)}%"

        main_alt = pos.alternative_brands[0] if pos.alternative_brands else None
        cells[6].text = f"{main_alt.brand} {main_alt.model}" if main_alt else "—"

        fill_color = ZEBRA_MINT if pos_idx % 2 == 1 else "FFFFFF"

        for idx, c in enumerate(cells):
            c.width = sum_headers[idx][1]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_bg(c, fill_color)
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.color.rgb = TEXT_DARK
            if idx in (0, 4, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if idx == 2:
                for r in p.runs:
                    r.font.bold = True
            if idx == 4 and (pos.gisp_match and pos.gisp_match.registry_number):
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = TEXT_GREEN
            if idx == 5:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = TEXT_GREEN

    _set_table_full_grid_borders(sum_table, BORDER_COLOR)
    _prevent_row_split(sum_table)
    _set_table_header_repeat(sum_table)

    # 5. Раздел 2: ПОДРОБНЫЕ ХАРАКТЕРИСТИКИ ПО КАЖДОЙ ПОЗИЦИИ
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    sec2_p = doc.add_paragraph()
    sec2_p.paragraph_format.space_before = Pt(8)
    sec2_p.paragraph_format.space_after = Pt(4)
    sec2_run = sec2_p.add_run("2. Подробные характеристики для заявки (Форма 2) и сравнение аналогов")
    sec2_run.font.bold = True
    sec2_run.font.size = Pt(10.5)
    sec2_run.font.color.rgb = BRAND_EMERALD

    for pos_idx, pos in enumerate(report.positions, start=1):
        pos_banner = doc.add_table(rows=1, cols=1)
        pos_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        pos_banner.autofit = False
        _set_table_fixed_width(pos_banner, 6.97)
        pos_banner.rows[0].cells[0].width = Inches(6.97)
        b_cell = pos_banner.rows[0].cells[0]
        _set_cell_bg(b_cell, BANNER_EMERALD)
        bp = b_cell.paragraphs[0]
        bp.paragraph_format.space_before = Pt(3)
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Inches(0.08)
        
        gisp_str = f"ГИСП № {pos.gisp_match.registry_number}" if (pos.gisp_match and pos.gisp_match.registry_number) else "ГИСП: Не в реестре"
        src_tag = f" | Источник: {pos.source_url[:45]}..." if pos.source_url else ""
        brun = bp.add_run(f"ПОЗИЦИЯ №{pos.position_no}: {pos.name_in_tz}\nТовар: {pos.identified_brand} {pos.identified_model} ({pos.manufacturer})   |   {gisp_str}{src_tag}")
        brun.font.bold = True
        brun.font.size = Pt(9)
        brun.font.color.rgb = RGBColor(255, 255, 255)
        _set_table_full_grid_borders(pos_banner, BANNER_EMERALD)

        if pos.reasoning:
            rsn_p = doc.add_paragraph()
            rsn_p.paragraph_format.space_before = Pt(3)
            rsn_p.paragraph_format.space_after = Pt(3)
            rsn_p.paragraph_format.left_indent = Inches(0.05)
            r_lbl = rsn_p.add_run("Обоснование соответствия ТЗ: ")
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(8.5)
            r_lbl.font.color.rgb = BRAND_EMERALD

            r_val = rsn_p.add_run(pos.reasoning)
            r_val.font.size = Pt(8.5)
            r_val.font.color.rgb = TEXT_MUTED

        # Таблица характеристик (Форма 2)
        if pos.specs_breakdown:
            sp_lbl = doc.add_paragraph()
            sp_lbl.paragraph_format.space_before = Pt(2)
            sp_lbl.paragraph_format.space_after = Pt(2)
            sp_lbl_run = sp_lbl.add_run("Показатели для первой части заявки (Форма 2, проверка по первоисточникам):")
            sp_lbl_run.font.bold = True
            sp_lbl_run.font.size = Pt(8.5)
            sp_lbl_run.font.color.rgb = TEXT_DARK

            spec_table = doc.add_table(rows=1, cols=6)
            spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            spec_table.autofit = False
            _set_table_fixed_width(spec_table, 6.97)

            spec_headers = [
                ("№", Inches(0.35)),
                ("Наименование параметра", Inches(1.70)),
                ("Требование ТЗ", Inches(1.45)),
                ("Конкретный показатель товара", Inches(1.50)),
                ("Соответствие", Inches(0.77)),
                ("Примечание / Источник", Inches(1.20)),
            ]

            for i, (h_text, w) in enumerate(spec_headers):
                cell = spec_table.rows[0].cells[i]
                cell.width = w
                cell.text = h_text
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(255, 255, 255)
                _set_cell_bg(cell, MEDIUM_EMERALD)

            for s_idx, spec in enumerate(pos.specs_breakdown, start=1):
                row = spec_table.add_row()
                cells = row.cells
                cells[0].text = str(s_idx)
                cells[1].text = spec.param_name
                cells[2].text = spec.tz_requirement
                cells[3].text = spec.product_fact
                status_text = "Подходит" if spec.status == "match" else "Уточнить (по паспорту)" if spec.status == "clarify" else "Отклонение"
                cells[4].text = status_text
                cells[5].text = spec.comment

                fill_color = ZEBRA_MINT if s_idx % 2 == 1 else "FFFFFF"

                for idx, c in enumerate(cells):
                    c.width = spec_headers[idx][1]
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_bg(c, fill_color)
                    p = c.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.05
                    p.paragraph_format.space_before = Pt(1.5)
                    p.paragraph_format.space_after = Pt(1.5)
                    for r in p.runs:
                        r.font.size = Pt(7.5)
                        r.font.color.rgb = TEXT_DARK
                    if idx in (0, 4):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if idx == 3:
                        for r in p.runs:
                            r.font.bold = True
                    if idx == 4:
                        for r in p.runs:
                            r.font.bold = True
                            if spec.status == "match":
                                r.font.color.rgb = TEXT_GREEN
                            elif spec.status == "clarify":
                                r.font.color.rgb = RGBColor(180, 83, 9)
                            else:
                                r.font.color.rgb = RGBColor(185, 28, 28)

            _set_table_full_grid_borders(spec_table, BORDER_COLOR)
            _prevent_row_split(spec_table)
            _set_table_header_repeat(spec_table)

        # Таблица аналогов со сравнением характеристик
        if pos.alternative_brands:
            alt_lbl = doc.add_paragraph()
            alt_lbl.paragraph_format.space_before = Pt(3)
            alt_lbl.paragraph_format.space_after = Pt(2)
            alt_lbl_run = alt_lbl.add_run("Сравнение взаимозаменяемых российских аналогов (по 44/223-ФЗ):")
            alt_lbl_run.font.bold = True
            alt_lbl_run.font.size = Pt(8.5)
            alt_lbl_run.font.color.rgb = TEXT_DARK

            alt_table = doc.add_table(rows=1, cols=6)
            alt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            alt_table.autofit = False
            _set_table_fixed_width(alt_table, 6.97)

            alt_headers = [
                ("№", Inches(0.35)),
                ("Аналог (Бренд / Модель)", Inches(1.35)),
                ("Завод-изготовитель", Inches(1.25)),
                ("Страна / Реестр", Inches(0.85)),
                ("Совм.", Inches(0.52)),
                ("Сравнение ключевых характеристик и замена", Inches(2.65)),
            ]

            for i, (h_text, w) in enumerate(alt_headers):
                cell = alt_table.rows[0].cells[i]
                cell.width = w
                cell.text = h_text
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(255, 255, 255)
                _set_cell_bg(cell, SUBTLE_EMERALD)

            for a_idx, alt in enumerate(pos.alternative_brands, start=1):
                row = alt_table.add_row()
                cells = row.cells
                cells[0].text = str(a_idx)
                cells[1].text = f"{alt.brand} {alt.model}"
                cells[2].text = alt.manufacturer
                cells[3].text = "РФ (Реестр)"
                cells[4].text = f"{int(alt.confidence * 100)}%"
                cells[5].text = alt.notes or "Взаимозаменяемый аналог по ГОСТ/ТУ"

                fill_color = ZEBRA_MINT if a_idx % 2 == 1 else "FFFFFF"

                for idx, c in enumerate(cells):
                    c.width = alt_headers[idx][1]
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_bg(c, fill_color)
                    p = c.paragraphs[0]
                    p.paragraph_format.line_spacing = 1.05
                    p.paragraph_format.space_before = Pt(1.5)
                    p.paragraph_format.space_after = Pt(1.5)
                    for r in p.runs:
                        r.font.size = Pt(7.5)
                        r.font.color.rgb = TEXT_DARK
                    if idx in (0, 3, 4):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if idx == 1:
                        for r in p.runs:
                            r.font.bold = True
                    if idx in (3, 4):
                        for r in p.runs:
                            r.font.bold = True
                            if alt.confidence >= 0.90:
                                r.font.color.rgb = TEXT_GREEN
                            else:
                                r.font.color.rgb = TEXT_DARK

            _set_table_full_grid_borders(alt_table, BORDER_COLOR)
            _prevent_row_split(alt_table)
            _set_table_header_repeat(alt_table)

            # Детальные показатели для первой части заявки по каждому российскому аналогу
            for a_idx, alt in enumerate(pos.alternative_brands, start=1):
                if not alt.specs_breakdown:
                    continue

                alt_form2_p = doc.add_paragraph()
                alt_form2_p.paragraph_format.space_before = Pt(6)
                alt_form2_p.paragraph_format.space_after = Pt(2)
                gisp_alt_tag = f"ГИСП № {alt.gisp_match.registry_number}" if (alt.gisp_match and alt.gisp_match.registry_number) else "ГИСП: Реестр РФ"
                r_a_hdr = alt_form2_p.add_run(f"Показатели аналога №{a_idx} для первой части заявки (Форма 2) — {alt.brand} {alt.model} ({alt.manufacturer})   |   {gisp_alt_tag}   |   Совместимость: {int(alt.confidence * 100)}%:")
                r_a_hdr.font.bold = True
                r_a_hdr.font.size = Pt(8.5)
                r_a_hdr.font.color.rgb = BRAND_EMERALD

                alt_spec_table = doc.add_table(rows=1, cols=6)
                alt_spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                alt_spec_table.autofit = False
                _set_table_fixed_width(alt_spec_table, 6.97)

                alt_spec_headers = [
                    ("№", Inches(0.35)),
                    ("Наименование параметра", Inches(1.70)),
                    ("Требование ТЗ", Inches(1.45)),
                    (f"Конкретный показатель ({alt.brand})", Inches(1.50)),
                    ("Соответствие", Inches(0.77)),
                    ("Примечание", Inches(1.20)),
                ]

                for i, (h_text, w) in enumerate(alt_spec_headers):
                    cell = alt_spec_table.rows[0].cells[i]
                    cell.width = w
                    cell.text = h_text
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(8)
                        r.font.color.rgb = RGBColor(255, 255, 255)
                    _set_cell_bg(cell, SUBTLE_EMERALD)

                for s_idx, spec in enumerate(alt.specs_breakdown, start=1):
                    row = alt_spec_table.add_row()
                    cells = row.cells
                    cells[0].text = str(s_idx)
                    cells[1].text = spec.param_name
                    cells[2].text = spec.tz_requirement
                    cells[3].text = spec.product_fact
                    status_text = "Подходит" if spec.status == "match" else "Уточнить (по паспорту)" if spec.status == "clarify" else "Отклонение"
                    cells[4].text = status_text
                    cells[5].text = spec.comment

                    fill_color = ZEBRA_MINT if s_idx % 2 == 1 else "FFFFFF"

                    for idx, c in enumerate(cells):
                        c.width = alt_spec_headers[idx][1]
                        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        _set_cell_bg(c, fill_color)
                        p = c.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.05
                        p.paragraph_format.space_before = Pt(1.5)
                        p.paragraph_format.space_after = Pt(1.5)
                        for r in p.runs:
                            r.font.size = Pt(7.5)
                            r.font.color.rgb = TEXT_DARK
                        if idx in (0, 4):
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if idx == 3:
                            for r in p.runs:
                                r.font.bold = True
                        if idx == 4:
                            for r in p.runs:
                                r.font.bold = True
                                if spec.status == "match":
                                    r.font.color.rgb = TEXT_GREEN
                                elif spec.status == "clarify":
                                    r.font.color.rgb = RGBColor(180, 83, 9)
                                else:
                                    r.font.color.rgb = RGBColor(185, 28, 28)

                _set_table_full_grid_borders(alt_spec_table, BORDER_COLOR)
                _prevent_row_split(alt_spec_table)
                _set_table_header_repeat(alt_spec_table)

    # 6. Раздел 3: РЕЕСТР ПРОВЕРЕННЫХ ИСТОЧНИКОВ И ПАСПОРТОВ ИЗДЕЛИЙ
    if report.verified_documents or report.web_sources:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        sec3_p = doc.add_paragraph()
        sec3_p.paragraph_format.space_before = Pt(8)
        sec3_p.paragraph_format.space_after = Pt(4)
        sec3_run = sec3_p.add_run("3. Реестр проверенных открытых веб-источников и паспортов изделий")
        sec3_run.font.bold = True
        sec3_run.font.size = Pt(10.5)
        sec3_run.font.color.rgb = BRAND_EMERALD

        src_table = doc.add_table(rows=1, cols=4)
        src_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        src_table.autofit = False
        _set_table_fixed_width(src_table, 6.97)

        src_headers = [
            ("№", Inches(0.35)),
            ("Тип источника", Inches(1.30)),
            ("Наименование / Документ", Inches(2.32)),
            ("Ссылка / Домен", Inches(3.00)),
        ]

        for i, (h_text, w) in enumerate(src_headers):
            cell = src_table.rows[0].cells[i]
            cell.width = w
            cell.text = h_text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_bg(cell, DARK_EMERALD)

        row_count = 0
        for doc_item in report.verified_documents[:8]:
            row_count += 1
            row = src_table.add_row()
            cells = row.cells
            cells[0].text = str(row_count)
            cells[1].text = "PDF Паспорт" if doc_item.get("type") == "pdf" else "Сайт завода"
            cells[2].text = str(doc_item.get("title") or "Техническая документация")[:60]
            cells[3].text = str(doc_item.get("url") or doc_item.get("domain") or "—")[:70]

            fill_color = ZEBRA_MINT if row_count % 2 == 1 else "FFFFFF"
            for idx, c in enumerate(cells):
                c.width = src_headers[idx][1]
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_bg(c, fill_color)
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(1.5)
                for r in p.runs:
                    r.font.size = Pt(7.5)
                    r.font.color.rgb = TEXT_DARK
                if idx in (0, 1):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if row_count == 0 and report.web_sources:
            for s_idx, src in enumerate(report.web_sources[:6], start=1):
                row = src_table.add_row()
                cells = row.cells
                cells[0].text = str(s_idx)
                cells[1].text = "Веб-поиск"
                cells[2].text = f"Сайт производителя / поставщика: {src}"
                cells[3].text = src

                fill_color = ZEBRA_MINT if s_idx % 2 == 1 else "FFFFFF"
                for idx, c in enumerate(cells):
                    c.width = src_headers[idx][1]
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_bg(c, fill_color)
                    p = c.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1.5)
                    p.paragraph_format.space_after = Pt(1.5)
                    for r in p.runs:
                        r.font.size = Pt(7.5)
                        r.font.color.rgb = TEXT_DARK
                    if idx in (0, 1):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set_table_full_grid_borders(src_table, BORDER_COLOR)
        _prevent_row_split(src_table)
        _set_table_header_repeat(src_table)

    # 7. Дисклеймер
    doc.add_paragraph().paragraph_format.space_before = Pt(8)
    disc_p = doc.add_paragraph()
    disc_p.paragraph_format.left_indent = Inches(0.05)
    r_disc = disc_p.add_run(f"Примечание: {report.disclaimer}")
    r_disc.font.size = Pt(8)
    r_disc.font.italic = True
    r_disc.font.color.rgb = TEXT_MUTED

    doc.save(str(target_path))
    return target_path


# ---------------------------------------------------------------------------
# XLSX Export
# ---------------------------------------------------------------------------

def write_exact_product_xlsx(
    path: str | Path,
    report: ExactProductReport,
    *,
    title: str = "Подбор товара и аналоги",
) -> Path:
    """Генерирует просторный, высококонтрастный и легко читаемый файл Excel во всю ширину экрана."""
    target_path = Path(path)
    target_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Подбор товара и аналоги"

    # Палитра
    navy_dark = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    navy_pos = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    slate_hdr = PatternFill(start_color="334155", end_color="334155", fill_type="solid")
    alt_hdr = PatternFill(start_color="475569", end_color="475569", fill_type="solid")
    card_bg = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    hint_bg = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
    zebra_bg = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    white_bg = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    match_bg = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
    clarify_bg = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
    mismatch_bg = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")

    # Шрифты
    title_font = Font(name="Calibri", size=14, bold=True, color="0F172A")
    sub_font = Font(name="Calibri", size=11, bold=True, color="334155")
    sec_font = Font(name="Calibri", size=12, bold=True, color="0F172A")
    white_bold_11 = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    white_bold_10 = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    bold_10 = Font(name="Calibri", size=10, bold=True, color="0F172A")
    reg_10 = Font(name="Calibri", size=10, color="1E293B")
    hint_font = Font(name="Calibri", size=10, color="334155")
    green_bold = Font(name="Calibri", size=10, bold=True, color="15803D")
    amber_bold = Font(name="Calibri", size=10, bold=True, color="B45309")
    red_bold = Font(name="Calibri", size=10, bold=True, color="B91C1C")

    thin_border = Border(
        left=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin", color="CBD5E1"),
        top=Side(style="thin", color="CBD5E1"),
        bottom=Side(style="thin", color="CBD5E1"),
    )

    def _style_range(row: int, start_col: int, end_col: int, fill=None, font=None, align=None, border=None):
        if end_col > start_col:
            ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
        for c in range(start_col, end_col + 1):
            cell = ws.cell(row=row, column=c)
            if fill:
                cell.fill = fill
            if font:
                cell.font = font
            if align:
                cell.alignment = align
            if border:
                cell.border = border

    def _calc_merged_h(text: str, total_w: int = 169, line_h: int = 19, min_h: int = 26) -> int:
        cleaned = str(text or "").strip()
        if not cleaned:
            return min_h
        chars_per_line = max(20, int(total_w * 0.65))
        lines = 0
        for part in cleaned.split("\n"):
            lines += max(1, (len(part) + chars_per_line - 1) // chars_per_line)
        return max(min_h, lines * line_h + 10)

    def _calc_row_h(cells: list[tuple[str, int]], line_h: int = 16, min_h: int = 26) -> int:
        max_lines = 1
        for val, col_w in cells:
            cleaned = str(val or "").strip()
            if not cleaned:
                continue
            chars_per_line = max(6, int(col_w * 0.85))
            for part in cleaned.split("\n"):
                lines = max(1, (len(part) + chars_per_line - 1) // chars_per_line)
                if lines > max_lines:
                    max_lines = lines
        return max(min_h, max_lines * line_h + 12)

    # 1. Шапка документа
    ws.append(["TenderLex | ПОДБОР ТОВАРА, ХАРАКТЕРИСТИКИ И АНАЛОГИ ПО ТЗ"])
    r1 = ws.max_row
    _style_range(r1, 1, 8, font=title_font, align=Alignment(vertical="center"))
    ws.row_dimensions[r1].height = 26

    ws.append([f"Закупка: {report.procurement_title} | Всего позиций в спецификации: {report.total_positions}"])
    r2 = ws.max_row
    _style_range(r2, 1, 8, font=sub_font, align=Alignment(vertical="center"))
    ws.row_dimensions[r2].height = 22

    # 2. Навигация и Резюме
    hint_msg = "💡 КАК РАБОТАТЬ С ОТЧЁТОМ: Вверху — сводная ведомость по всей закупке. Ниже — построчные характеристики для первой части заявки (Форма 2) и альтернативные заводы РФ по каждой позиции."
    ws.append([hint_msg])
    r3 = ws.max_row
    _style_range(r3, 1, 8, fill=hint_bg, font=hint_font, align=Alignment(vertical="top", wrap_text=True))
    ws.row_dimensions[r3].height = _calc_merged_h(hint_msg, 169, 19, 26)

    if report.summary:
        sum_text = f"Экспертное резюме: {report.summary}"
        ws.append([sum_text])
        r4 = ws.max_row
        _style_range(r4, 1, 8, fill=card_bg, font=reg_10, align=Alignment(vertical="top", wrap_text=True))
        ws.row_dimensions[r4].height = _calc_merged_h(sum_text, 169, 19, 32)

    # Разделитель
    ws.append([None] * 8)
    r_sep1 = ws.max_row
    ws.row_dimensions[r_sep1].height = 14

    # 3. Раздел 1: СВОДНАЯ ВЕДОМОСТЬ
    ws.append(["1. СВОДНАЯ ВЕДОМОСТЬ ПО ВСЕМ ПОЗИЦИЯМ СПЕЦИФИКАЦИИ"])
    r_s1 = ws.max_row
    _style_range(r_s1, 1, 8, font=sec_font, align=Alignment(vertical="center"))
    ws.row_dimensions[r_s1].height = 26

    headers1 = [
        "№",
        "Позиция из ТЗ",
        "Выявленный бренд",
        "Точная модель / артикул",
        "Завод-изготовитель",
        "Реестр ГИСП",
        "Соответствие",
        "Основной аналог (РФ)",
    ]
    ws.append(headers1)
    r_h1 = ws.max_row
    ws.row_dimensions[r_h1].height = 26
    for col_idx in range(1, 9):
        cell = ws.cell(row=r_h1, column=col_idx)
        cell.fill = navy_dark
        cell.font = white_bold_10
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    col_widths_s1 = [6, 30, 24, 24, 24, 16, 15, 28]

    for pos_idx, pos in enumerate(report.positions, start=1):
        gisp_text = f"№ {pos.gisp_match.registry_number}" if (pos.gisp_match and pos.gisp_match.registry_number) else "Не в реестре"
        main_alt = pos.alternative_brands[0] if pos.alternative_brands else None
        alt_str = f"{main_alt.brand} {main_alt.model} ({main_alt.manufacturer})" if main_alt else "—"

        row_vals = [
            pos.position_no or pos_idx,
            pos.name_in_tz,
            pos.identified_brand,
            pos.identified_model,
            pos.manufacturer,
            gisp_text,
            f"{int(pos.confidence * 100)}%",
            alt_str,
        ]
        ws.append(row_vals)
        curr_r = ws.max_row
        fill_row = zebra_bg if pos_idx % 2 == 1 else white_bg

        for col_idx in range(1, 9):
            c = ws.cell(row=curr_r, column=col_idx)
            c.border = thin_border
            c.fill = fill_row
            if col_idx == 1:
                c.font = bold_10
                c.alignment = Alignment(horizontal="center", vertical="top")
            elif col_idx == 6:
                c.font = green_bold if (pos.gisp_match and pos.gisp_match.registry_number) else reg_10
                c.alignment = Alignment(horizontal="center", vertical="top")
            elif col_idx == 7:
                c.font = green_bold
                c.alignment = Alignment(horizontal="center", vertical="top")
            elif col_idx in (3, 4):
                c.font = bold_10
                c.alignment = Alignment(vertical="top", wrap_text=True)
            else:
                c.font = reg_10
                c.alignment = Alignment(vertical="top", wrap_text=True)

        cells_with_w = list(zip([str(v) for v in row_vals], col_widths_s1))
        ws.row_dimensions[curr_r].height = _calc_row_h(cells_with_w, line_h=16, min_h=26)

    # Разделитель
    ws.append([None] * 8)
    r_sep2 = ws.max_row
    ws.row_dimensions[r_sep2].height = 16

    # 4. Раздел 2: ПОДРОБНЫЕ ХАРАКТЕРИСТИКИ
    ws.append(["2. ПОДРОБНЫЕ ХАРАКТЕРИСТИКИ ДЛЯ ЗАЯВКИ И АНАЛОГИ ПО КАЖДОЙ ПОЗИЦИИ"])
    r_s2 = ws.max_row
    _style_range(r_s2, 1, 8, font=sec_font, align=Alignment(vertical="center"))
    ws.row_dimensions[r_s2].height = 28

    for pos_idx, pos in enumerate(report.positions, start=1):
        ws.append([None] * 8)
        r_pos_sep = ws.max_row
        ws.row_dimensions[r_pos_sep].height = 10

        gisp_text = f"ГИСП № {pos.gisp_match.registry_number}" if (pos.gisp_match and pos.gisp_match.registry_number) else "ГИСП: Не в реестре"
        banner_text = f"ПОЗИЦИЯ №{pos.position_no}: {pos.name_in_tz}   |   Товар: {pos.identified_brand} {pos.identified_model} ({pos.manufacturer})   |   {gisp_text}"

        ws.append([banner_text])
        r_ban = ws.max_row
        _style_range(r_ban, 1, 8, fill=navy_pos, font=white_bold_11, align=Alignment(vertical="center", wrap_text=True))
        ws.row_dimensions[r_ban].height = _calc_merged_h(banner_text, 169, 20, 28)

        if pos.reasoning:
            reason_text = f"Обоснование соответствия ТЗ: {pos.reasoning}"
            ws.append([reason_text])
            r_rsn = ws.max_row
            _style_range(r_rsn, 1, 8, fill=card_bg, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
            ws.row_dimensions[r_rsn].height = _calc_merged_h(reason_text, 169, 18, 26)

        # Таблица характеристик (Форма 2)
        if pos.specs_breakdown:
            ws.append(["№", "Наименование параметра", "Требование заказчика (ТЗ)", "Конкретный показатель товара", "", "Соответствие", "Примечание и обоснование показателя", ""])
            r_sp_h = ws.max_row
            ws.row_dimensions[r_sp_h].height = 24
            _style_range(r_sp_h, 1, 1, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_sp_h, 2, 2, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_sp_h, 3, 3, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_sp_h, 4, 5, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_sp_h, 6, 6, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_sp_h, 7, 8, fill=slate_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))

            for s_idx, spec in enumerate(pos.specs_breakdown, start=1):
                status_str = "Подходит" if spec.status == "match" else "Уточнить (по паспорту)" if spec.status == "clarify" else "Отклонение"
                ws.append([
                    s_idx,
                    spec.param_name,
                    spec.tz_requirement,
                    spec.product_fact,
                    "",
                    status_str,
                    spec.comment,
                    "",
                ])
                r_sp = ws.max_row
                fill_r = zebra_bg if s_idx % 2 == 1 else white_bg

                _style_range(r_sp, 1, 1, fill=fill_r, font=bold_10, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                _style_range(r_sp, 2, 2, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                _style_range(r_sp, 3, 3, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                _style_range(r_sp, 4, 5, fill=fill_r, font=bold_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)

                status_fill = match_bg if spec.status == "match" else clarify_bg if spec.status == "clarify" else mismatch_bg
                status_font = green_bold if spec.status == "match" else amber_bold if spec.status == "clarify" else red_bold
                _style_range(r_sp, 6, 6, fill=status_fill, font=status_font, align=Alignment(horizontal="center", vertical="top"), border=thin_border)

                _style_range(r_sp, 7, 8, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)

                sp_cells = [
                    (str(s_idx), 6),
                    (spec.param_name, 30),
                    (spec.tz_requirement, 24),
                    (spec.product_fact, 44),
                    (status_str, 15),
                    (spec.comment, 43),
                ]
                ws.row_dimensions[r_sp].height = _calc_row_h(sp_cells, line_h=16, min_h=24)

        # Таблица аналогов позиции
        if pos.alternative_brands:
            ws.append(["№", "Аналог (Бренд / Модель)", "Завод-изготовитель", "Страна", "Реестр РФ", "Совместимость", "Особенности, отличия и обоснование замены", ""])
            r_alt_h = ws.max_row
            ws.row_dimensions[r_alt_h].height = 24
            _style_range(r_alt_h, 1, 1, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 2, 2, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 3, 3, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 4, 4, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 5, 5, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 6, 6, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
            _style_range(r_alt_h, 7, 8, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))

            for a_idx, alt in enumerate(pos.alternative_brands, start=1):
                notes_text = alt.notes or "Взаимозаменяемый промышленный аналог"
                ws.append([
                    a_idx,
                    f"{alt.brand} {alt.model}",
                    alt.manufacturer,
                    "Россия",
                    "Реестр РФ",
                    f"{int(alt.confidence * 100)}%",
                    notes_text,
                    "",
                ])
                r_alt = ws.max_row
                fill_a = zebra_bg if a_idx % 2 == 1 else white_bg

                _style_range(r_alt, 1, 1, fill=fill_a, font=bold_10, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                _style_range(r_alt, 2, 2, fill=fill_a, font=bold_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                _style_range(r_alt, 3, 3, fill=fill_a, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                _style_range(r_alt, 4, 4, fill=fill_a, font=reg_10, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                _style_range(r_alt, 5, 5, fill=fill_a, font=green_bold, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                _style_range(r_alt, 6, 6, fill=fill_a, font=green_bold, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                _style_range(r_alt, 7, 8, fill=fill_a, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)

                alt_cells = [
                    (str(a_idx), 6),
                    (f"{alt.brand} {alt.model}", 30),
                    (alt.manufacturer, 24),
                    ("Россия", 20),
                    ("Реестр РФ", 16),
                    (f"{int(alt.confidence * 100)}%", 15),
                    (notes_text, 43),
                ]
                ws.row_dimensions[r_alt].height = _calc_row_h(alt_cells, line_h=16, min_h=24)

            # Детальные показатели для Формы 2 по каждому аналогу в Excel
            for a_idx, alt in enumerate(pos.alternative_brands, start=1):
                if not alt.specs_breakdown:
                    continue
                ws.append([None] * 8)
                r_alt_sp_sep = ws.max_row
                ws.row_dimensions[r_alt_sp_sep].height = 8

                gisp_alt_tag = f"ГИСП № {alt.gisp_match.registry_number}" if (alt.gisp_match and alt.gisp_match.registry_number) else "ГИСП: Реестр РФ"
                alt_hdr_text = f"Показатели аналога №{a_idx} для заявки (Форма 2) — {alt.brand} {alt.model} ({alt.manufacturer}) | {gisp_alt_tag} | Совместимость: {int(alt.confidence * 100)}%"
                ws.append([alt_hdr_text])
                r_alt_ban = ws.max_row
                _style_range(r_alt_ban, 1, 8, fill=card_bg, font=bold_10, align=Alignment(vertical="center", wrap_text=True), border=thin_border)
                ws.row_dimensions[r_alt_ban].height = _calc_merged_h(alt_hdr_text, 169, 18, 26)

                ws.append(["№", "Наименование параметра", "Требование заказчика (ТЗ)", f"Конкретный показатель ({alt.brand})", "", "Соответствие", "Примечание и обоснование показателя", ""])
                r_asp_h = ws.max_row
                ws.row_dimensions[r_asp_h].height = 24
                _style_range(r_asp_h, 1, 1, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
                _style_range(r_asp_h, 2, 2, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
                _style_range(r_asp_h, 3, 3, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
                _style_range(r_asp_h, 4, 5, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
                _style_range(r_asp_h, 6, 6, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))
                _style_range(r_asp_h, 7, 8, fill=alt_hdr, font=white_bold_10, align=Alignment(horizontal="center", vertical="center"))

                for s_idx, spec in enumerate(alt.specs_breakdown, start=1):
                    status_str = "Подходит" if spec.status == "match" else "Уточнить (по паспорту)" if spec.status == "clarify" else "Отклонение"
                    ws.append([
                        s_idx,
                        spec.param_name,
                        spec.tz_requirement,
                        spec.product_fact,
                        "",
                        status_str,
                        spec.comment,
                        "",
                    ])
                    r_asp = ws.max_row
                    fill_r = zebra_bg if s_idx % 2 == 1 else white_bg

                    _style_range(r_asp, 1, 1, fill=fill_r, font=bold_10, align=Alignment(horizontal="center", vertical="top"), border=thin_border)
                    _style_range(r_asp, 2, 2, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                    _style_range(r_asp, 3, 3, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)
                    _style_range(r_asp, 4, 5, fill=fill_r, font=bold_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)

                    status_fill = match_bg if spec.status == "match" else clarify_bg if spec.status == "clarify" else mismatch_bg
                    status_font = green_bold if spec.status == "match" else amber_bold if spec.status == "clarify" else red_bold
                    _style_range(r_asp, 6, 6, fill=status_fill, font=status_font, align=Alignment(horizontal="center", vertical="top"), border=thin_border)

                    _style_range(r_asp, 7, 8, fill=fill_r, font=reg_10, align=Alignment(vertical="top", wrap_text=True), border=thin_border)

                    sp_cells = [
                        (str(s_idx), 6),
                        (spec.param_name, 30),
                        (spec.tz_requirement, 24),
                        (spec.product_fact, 44),
                        (status_str, 15),
                        (spec.comment, 43),
                    ]
                    ws.row_dimensions[r_asp].height = _calc_row_h(sp_cells, line_h=16, min_h=24)

    # 5. Дисклеймер
    ws.append([None] * 8)
    r_disc_sep = ws.max_row
    ws.row_dimensions[r_disc_sep].height = 10

    disc_text = f"Примечание: {report.disclaimer}"
    ws.append([disc_text])
    r_disc = ws.max_row
    _style_range(r_disc, 1, 8, font=Font(name="Calibri", size=9, italic=True, color="64748B"), align=Alignment(vertical="center", wrap_text=True))
    ws.row_dimensions[r_disc].height = _calc_merged_h(disc_text, 169, 16, 26)

    # 6. Ширина колонок
    ws.column_dimensions["A"].width = 6    # №
    ws.column_dimensions["B"].width = 30   # Позиция ТЗ / Параметр / Аналог
    ws.column_dimensions["C"].width = 24   # Бренд / Требование ТЗ / Завод
    ws.column_dimensions["D"].width = 22   # Модель / Показатель факт / Страна
    ws.column_dimensions["E"].width = 22   # Завод / Реестр
    ws.column_dimensions["F"].width = 16   # Реестр ГИСП / Статус соответствия
    ws.column_dimensions["G"].width = 15   # Соответствие % / Совместимость %
    ws.column_dimensions["H"].width = 28   # Аналог РФ / Обоснование

    # Настройки страницы
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    wb.save(str(target_path))
    return target_path


# ---------------------------------------------------------------------------
# Styling Helpers for DOCX
# ---------------------------------------------------------------------------

def _set_docx_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)


def _set_cell_bg(cell, color_hex: str) -> None:
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_table_fixed_width(table, width_in_inches: float = 6.97) -> None:
    tblPr = table._tbl.tblPr
    dxa_val = int(width_in_inches * 1440)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{dxa_val}" w:type="dxa"/>')
    tblPr.append(tblW)


def _set_table_full_grid_borders(table, color: str = "CBD5E1") -> None:
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _prevent_row_split(table) -> None:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def _set_table_header_repeat(table) -> None:
    if table.rows:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

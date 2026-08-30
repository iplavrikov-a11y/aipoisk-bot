from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from openpyxl import load_workbook

from app.quote_request import build_quote_request_markdown
from app.report_builder import (
    PROCUREMENT_REPORT_DISCLAIMER,
    QUOTE_REQUEST_INTRO,
    REGISTRY_FALLBACK_REPORT_DISCLAIMER,
    write_procurement_docx,
    write_quote_request_docx,
    write_supplier_xlsx,
)


class ReportBuilderTests(unittest.TestCase):
    def test_procurement_docx_contains_soft_ai_disclaimer(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "analysis.docx"
            write_procurement_docx(
                path,
                "#### Общая информация\n- Заказчик: Тестовый заказчик",
                title="Анализ документации",
            )

            doc = Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

            self.assertIn(PROCUREMENT_REPORT_DISCLAIMER, text)
            self.assertIn("Критичные юридические, финансовые и технические условия", text)

    def test_procurement_docx_heading_and_body_font_sizes(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "analysis.docx"
            write_procurement_docx(
                path,
                "#### Общая информация\n- Заказчик: Тестовый заказчик\n- ИНН: 7701234567",
                title="Анализ документации",
            )

            doc = Document(path)
            heading_p = [p for p in doc.paragraphs if p.text == "Общая информация"][0]
            body_p = [p for p in doc.paragraphs if "Тестовый заказчик" in p.text][0]

            heading_size = heading_p.runs[0].font.size.pt
            body_size = body_p.runs[0].font.size.pt

            self.assertEqual(heading_size, 11.0)
            self.assertEqual(body_size, 9.5)
            self.assertGreater(heading_size, body_size)

    def test_procurement_docx_removes_okpd_and_formats_tz_columns(self) -> None:
        from docx import Document

        markdown = """#### Товары и требования (Техническое задание)
| № | Наименование | Характеристики | Ед.изм. | Кол-во |
|---|---|---|---|---|
| 1 | Канат стальной (ОКПД2: 25.93.11.120) | Назначение: применяется в качестве устройств растяжек и вант для кранов; требования к конструкции: ГОСТ 3062-80, конструкция 1x7, диаметр 6,8 мм. | м | 5000 |
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "analysis.docx"
            write_procurement_docx(path, markdown, title="Анализ закупки")

            doc = Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)
            with zipfile.ZipFile(path) as archive:
                xml = archive.read("word/document.xml")

        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        grid_cols = [
            int(col.attrib[f"{{{ns['w']}}}w"])
            for col in root.findall(".//w:tbl[1]/w:tblGrid/w:gridCol", ns)
        ]

        self.assertNotRegex(f"{text}\n{table_text}", r"(?i)ОКПД|OKPD")
        self.assertNotIn("25.93.11.120", table_text)
        self.assertEqual(len(grid_cols), 5)
        self.assertLess(grid_cols[0], grid_cols[1])
        self.assertGreater(grid_cols[2], grid_cols[1])
        self.assertGreater(grid_cols[2], grid_cols[3] * 5)
        self.assertGreater(grid_cols[2], grid_cols[4] * 5)

    def test_quote_request_docx_uses_request_text_without_analysis_disclaimer(self) -> None:
        from docx import Document

        markdown = f"""ЗАПРОС КП

{QUOTE_REQUEST_INTRO}

| № | Наименование | Характеристики | Ед.изм. | Кол-во |
|---|---|---|---|---|
| 1 | Канат стальной КТРУ 25.93.11 | ГОСТ 3062-80, диаметр 6,8 мм | м | 5000 |

### Условия поставки

- **Срок поставки:** 30 календарных дней
- **Город поставки:** Казань
- **Документы качества:** паспорт качества
- **Упаковка/тара:** бухты
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "quote.docx"
            write_quote_request_docx(path, markdown, title="Запрос КП")

            doc = Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for row in doc.tables[0].rows for cell in row.cells)

        self.assertIn("Запрос КП", text)
        self.assertIn(QUOTE_REQUEST_INTRO, text)
        self.assertIn("Канат стальной", table_text)
        self.assertIn("Срок поставки:", text)
        self.assertNotIn(PROCUREMENT_REPORT_DISCLAIMER, text)
        self.assertNotRegex(f"{text}\n{table_text}", r"(?i)ОКПД|OKPD|КТРУ|KTRU")

    def test_quote_request_markdown_omits_note_column(self) -> None:
        source = """### Товары и требования
| № | Наименование | Характеристики | Ед.изм. | Кол-во | Примечание |
|---|---|---|---|---|---|
| 1 | Канат стальной | ГОСТ 3062-80, диаметр 6,8 мм | м | 5000 | Просим указать в КП |

- **Срок поставки:** 30 календарных дней
- **Город поставки:** Казань
- **Условия оплаты:** по договору
- **Документы качества:** паспорт качества
- **Упаковка/тара:** бухты
"""

        markdown = build_quote_request_markdown(source, subject="Канат стальной")

        self.assertIn("| № | Наименование | Характеристики | Ед.изм. | Кол-во |", markdown)
        self.assertNotIn("Примечание", markdown)
        self.assertNotIn("Просим указать в КП", markdown)

    def test_supplier_xlsx_is_client_facing_without_technical_search_columns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Поставщик",
                        "match_level": "exact",
                        "quality_score": 91,
                        "quality_tier": "high",
                        "procurement_item": "Сварочный полуавтомат",
                        "ai_confidence": 92,
                        "site_type": "manufacturer",
                        "product_fit": "exact",
                        "product": "Сварочный полуавтомат MIG/MAG",
                        "contact_person": "Отдел продаж",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "evidence_snippet": "В каталоге указан сварочный полуавтомат.",
                        "contact_evidence_snippet": "Контакты: +7 999 111 22 33",
                        "source": "yandex",
                        "contact_url": "https://supplier.ru/contacts",
                        "evidence_url": "https://supplier.ru/catalog",
                    }
                ],
                title="ТЗ",
                subject="Сварочный полуавтомат",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            sheet_title = ws.title
            report_title = ws["A1"].value
            headers = [cell.value for cell in ws[5]]
            values = [cell.value for cell in ws[6]]
            site_hyperlink = ws["B6"].hyperlink.target if ws["B6"].hyperlink else ""
            self.assertIn("Запрос КП", wb.sheetnames)
            ws_quote = wb["Запрос КП"]
            quote_text = ws_quote["A3"].value
            self.assertIsNone(ws.freeze_panes)
            self.assertFalse(ws.auto_filter.ref)
            wb.close()

        self.assertEqual(sheet_title, "Поставщики")
        self.assertIn("TenderLex", report_title)
        self.assertIn("Сварочный полуавтомат", report_title)
        self.assertIn("Тема письма: Запрос коммерческого предложения", quote_text)
        self.assertIn("Добрый день!\n\n", quote_text)
        self.assertNotIn("отдел продаж", quote_text.lower())
        self.assertIn("С уважением", quote_text)
        self.assertEqual(
            headers,
            [
                "Компания",
                "Сайт",
                "Телефоны",
                "Email",
                "Комментарий",
                "Реестр Минпромторга",
            ],
        )
        self.assertNotIn("Соответствие", headers)
        self.assertNotIn("Search Query", headers)
        self.assertNotIn("Evidence snippet", headers)
        self.assertNotIn("AI уверенность", headers)
        self.assertNotIn("Contact URL", headers)
        self.assertIn("Поставщик", values)
        self.assertIn("https://supplier.ru", values)
        self.assertIn("+7 999 111 22 33", values)
        self.assertIn("sales@supplier.ru", values)
        self.assertEqual(site_hyperlink, "https://supplier.ru")

    def test_supplier_xlsx_keeps_font_children_in_schema_order(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [{"company_name": "Поставщик", "product_fit": "exact"}],
                title="ТЗ",
                target=1,
            )
            with zipfile.ZipFile(path) as archive:
                styles = ET.fromstring(archive.read("xl/styles.xml"))

        namespace = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
        order = (
            "b", "i", "strike", "outline", "shadow", "condense", "extend", "sz", "u",
            "vertAlign", "color", "name", "charset", "family", "scheme",
        )
        positions = {name: index for index, name in enumerate(order)}
        for font in styles.findall(f".//{{{namespace}}}fonts/{{{namespace}}}font"):
            child_positions = [
                positions[child.tag.rsplit("}", 1)[-1]]
                for child in font
                if child.tag.rsplit("}", 1)[-1] in positions
            ]
            self.assertEqual(child_positions, sorted(child_positions))

    def test_supplier_xlsx_marks_non_exact_matches_as_confirmation_needed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Профильный поставщик",
                        "product_fit": "profile",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "comments": "Компания полностью соответствует ТЗ по профилю закупки.",
                    }
                ],
                title="ТЗ",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            comment = ws["E6"].value
            wb.close()

        self.assertIn("Профиль компании подходит.", comment)
        self.assertIn("Наличие товара уточнить", comment)
        self.assertNotIn("полностью соответствует", comment.lower())

    def test_supplier_xlsx_uses_short_exact_comment(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Поставщик",
                        "product_fit": "exact",
                        "product": "Сварочный полуавтомат MIG/MAG 500А с блоком охлаждения и комплектом поставки",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "comments": "Очень длинное описание, которое не должно попадать в клиентский отчет целиком.",
                    }
                ],
                title="ТЗ",
                subject="Сварочный полуавтомат",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            comment = ws["E6"].value
            wb.close()

        self.assertIn("Точное соответствие:", comment)
        self.assertLessEqual(len(comment), 260)

    def test_supplier_xlsx_does_not_expose_minprom_notes_in_comment(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Поставщик",
                        "product_fit": "category",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "comments": (
                            "Нужно запросить подтверждение реестровой записи Минпромторга. "
                            "Поставщик релевантен по промышленной категории."
                        ),
                    }
                ],
                title="ТЗ",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            comment = ws["E6"].value
            wb.close()

        self.assertIn("Категория совпадает", comment)
        self.assertNotRegex(comment, r"(?i)минпромторг|гисп|реестров")

    def test_supplier_xlsx_adds_registry_fallback_note_for_priority_mode(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Поставщик",
                        "product_fit": "exact",
                        "product": "Компаратор видеоспектральный Regula 4308",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "supplier_search_policy": "minprom_registry_priority",
                        "supplier_search_origin": "ordinary_fallback",
                        "minprom_registry_required": True,
                        "minprom_registry_status": "empty",
                        "minprom_registry_match": {"matched": False},
                    }
                ],
                title="ТЗ",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            comment = ws["E6"].value
            wb.close()

        self.assertIn("Точное соответствие", comment)
        self.assertIn("Реестр: релевантная запись не найдена", comment)
        self.assertIn("обычным поиском", comment)

    def test_supplier_xlsx_marks_strict_registry_fallback_at_document_and_row_level(self) -> None:
        for registry_status in ("empty", "ok"):
            with self.subTest(registry_status=registry_status), tempfile.TemporaryDirectory() as tmp:
                path = Path(tmp) / "suppliers-without-registry-confirmation.xlsx"
                write_supplier_xlsx(
                    path,
                    [
                        {
                            "company_name": "Поставщик",
                            "product_fit": "exact",
                            "product": "Волоконный усилитель",
                            "phone": "+7 999 111 22 33",
                            "email": "sales@supplier.ru",
                            "site": "https://supplier.ru",
                            "supplier_search_policy": "minprom_registry_only",
                            "supplier_search_origin": "ordinary_fallback",
                            "minprom_registry_required": True,
                            "minprom_registry_status": registry_status,
                            "minprom_registry_match": {"matched": False},
                        }
                    ],
                    title="ТЗ",
                    target=1,
                )

                wb = load_workbook(path)
                ws = wb["Поставщики"]
                document_warning = ws["A3"].value
                row_comment = ws["E6"].value
                warning_fill = ws["A3"].fill.fgColor.rgb
                wb.close()

            self.assertIn(REGISTRY_FALLBACK_REPORT_DISCLAIMER, document_warning)
            self.assertIn("не подтверждено", document_warning)
            self.assertNotIn("запись подтверждена", document_warning.lower())
            self.assertIn("Точное соответствие", row_comment)
            self.assertIn("Реестр:", row_comment)
            self.assertNotIn("запись подтверждена", row_comment.lower())
            if registry_status == "empty":
                self.assertIn("релевантная запись не найдена", row_comment)
            else:
                self.assertIn("соответствие поставщика конкретной записи не подтверждено", row_comment)
            self.assertEqual(warning_fill, "00FEF3C7")

    def test_supplier_xlsx_adds_registry_number_in_single_comment_column(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Завод",
                        "product_fit": "exact",
                        "product": "Насос центробежный",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                        "supplier_search_policy": "minprom_registry_only",
                        "supplier_search_origin": "minprom_registry",
                        "minprom_registry_required": True,
                        "minprom_registry_status": "ok",
                        "minprom_registry_match": {
                            "matched": True,
                            "registry_number": "РПП-123",
                            "manufacturer": 'АО "Завод"',
                        },
                    }
                ],
                title="ТЗ",
                target=1,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            headers = [cell.value for cell in ws[5]]
            comment = ws["E6"].value
            registry_cell = ws["F6"].value
            wb.close()

        self.assertEqual(
            headers,
            [
                "Компания",
                "Сайт",
                "Телефоны",
                "Email",
                "Комментарий",
                "Реестр Минпромторга",
            ],
        )
        self.assertIn("Реестр ГИСП Минпромторга: запись № РПП-123", comment)
        self.assertIn('АО "Завод"', comment)
        self.assertIn("РПП-123", registry_cell)

    def test_supplier_xlsx_hides_internal_target_when_overfilled(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": f"Поставщик {index}",
                        "product_fit": "category",
                        "product": "Стальные канаты",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": f"https://supplier-{index}.ru",
                    }
                    for index in range(4)
                ],
                title="ТЗ",
                target=2,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            summary = ws["A3"].value
            wb.close()

        self.assertIn("Проверены сайты и контакты. Кандидатов: 4", summary)
        self.assertIn("точным техническим совпадением: 0", summary)
        self.assertIn("категорийных кандидатов: 4", summary)
        self.assertNotIn("Минимум по настройкам", summary)
        self.assertNotIn("4/2", summary)

    def test_supplier_xlsx_hides_internal_target_when_underfilled(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "suppliers.xlsx"
            write_supplier_xlsx(
                path,
                [
                    {
                        "company_name": "Поставщик",
                        "product_fit": "category",
                        "product": "Промышленное оборудование",
                        "phone": "+7 999 111 22 33",
                        "email": "sales@supplier.ru",
                        "site": "https://supplier.ru",
                    }
                ],
                title="ТЗ",
                target=15,
            )

            wb = load_workbook(path)
            ws = wb["Поставщики"]
            summary = ws["A3"].value
            wb.close()

        self.assertIn("Проверены сайты и контакты. Кандидатов: 1", summary)
        self.assertIn("точным техническим совпадением: 0", summary)
        self.assertIn("категорийных кандидатов: 1", summary)
        self.assertNotIn("1/15", summary)
        self.assertNotIn("миним", summary.lower())


if __name__ == "__main__":
    unittest.main()

from __future__ import annotations

import csv
import html
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from bs4 import BeautifulSoup


TEXT_EXTENSIONS = {".txt", ".md", ".json", ".yaml", ".yml", ".ini", ".log"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}
DEFAULT_DOCUMENT_OPTIONS = {
    "ocr_enabled": True,
    "pdf_ocr_pages": 3,
    "archive_max_files": 80,
    "archive_max_mb": 250,
    "archive_depth": 2,
}


def sanitize_filename(value: str) -> str:
    value = re.sub(r"[^\wА-Яа-яЁё ._-]+", "_", str(value or ""), flags=re.UNICODE).strip(" ._")
    return value[:160] or "upload"


def extract_text(path: str | Path, options: dict | None = None, _depth: int | None = None) -> tuple[str, str]:
    options = {**DEFAULT_DOCUMENT_OPTIONS, **(options or {})}
    depth = int(options.get("archive_depth") or 0) if _depth is None else _depth
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    try:
        if suffix in TEXT_EXTENSIONS:
            return file_path.read_text(encoding="utf-8", errors="ignore"), "ok"
        if suffix in {".csv", ".tsv"}:
            return _extract_csv(file_path), "ok"
        if suffix in {".html", ".htm", ".xml"}:
            return _extract_html(file_path), "ok"
        if suffix == ".docx":
            return _extract_docx(file_path), "ok"
        if suffix == ".xlsx":
            return _extract_xlsx(file_path), "ok"
        if suffix == ".xls":
            text = _extract_via_libreoffice(file_path, ".xlsx", _extract_xlsx)
            return text, "ok" if text.strip() else "parser_not_connected_yet"
        if suffix == ".pdf":
            text = _extract_pdf(file_path)
            if len(text.strip()) < 80 and options.get("ocr_enabled"):
                ocr_text = _extract_pdf_ocr(file_path, int(options.get("pdf_ocr_pages") or 0))
                if ocr_text.strip():
                    return ocr_text, "pdf_ocr_ok"
            return text, "ok"
        if suffix in IMAGE_EXTENSIONS:
            if options.get("ocr_enabled"):
                text = _extract_image_ocr(file_path)
                return text, "image_ocr_ok" if text.strip() else "image_ocr_empty"
            return "", "image_ocr_disabled"
        if suffix == ".doc":
            text = _extract_doc(file_path)
            return text, "ok" if text.strip() else "parser_not_connected_yet"
        if suffix in {".rtf", ".odt", ".pptx"}:
            text = _extract_via_pandoc(file_path) or _extract_via_libreoffice(
                file_path,
                ".txt",
                lambda converted: converted.read_text(encoding="utf-8", errors="ignore"),
            )
            return text, "ok" if text.strip() else "parser_not_connected_yet"
        if suffix in ARCHIVE_EXTENSIONS:
            if depth <= 0:
                return "", "archive_depth_limit"
            return _extract_archive(file_path, options, depth)
        return "", "unsupported_extension"
    except Exception as exc:
        return "", f"error:{exc.__class__.__name__}: {exc}"


def _extract_csv(path: Path) -> str:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    rows = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        for row in csv.reader(handle, delimiter=delimiter):
            rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def _extract_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return html.unescape(soup.get_text("\n", strip=True))


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    blocks: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        blocks.append(f"\n=== TABLE {table_index} ===")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    blocks: list[str] = []
    for ws in wb.worksheets:
        blocks.append(f"\n=== SHEET: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if value is None else str(value).strip() for value in row]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _extract_pdf(path: Path) -> str:
    try:
        import fitz

        doc = fitz.open(str(path))
        return "\n".join(page.get_text("text") for page in doc)
    except Exception:
        import PyPDF2

        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_doc(path: Path) -> str:
    antiword = shutil.which("antiword")
    if antiword:
        result = subprocess.run([antiword, str(path)], check=False, capture_output=True, text=True, timeout=60)
        if result.stdout.strip():
            return result.stdout
    return _extract_via_libreoffice(
        path,
        ".txt",
        lambda converted: converted.read_text(encoding="utf-8", errors="ignore"),
    )


def _extract_via_pandoc(path: Path) -> str:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return ""
    result = subprocess.run([pandoc, str(path), "-t", "plain"], check=False, capture_output=True, text=True, timeout=80)
    return result.stdout if result.returncode == 0 else ""


def _extract_via_libreoffice(path: Path, target_suffix: str, reader) -> str:
    libreoffice = shutil.which("libreoffice")
    if not libreoffice:
        return ""
    with tempfile.TemporaryDirectory(prefix="aipoisk-lo-") as tmp:
        result = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to",
                target_suffix.lstrip("."),
                "--outdir",
                tmp,
                str(path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            return ""
        candidates = sorted(Path(tmp).glob(f"*{target_suffix}"))
        if not candidates:
            return ""
        return reader(candidates[0])


def _extract_image_ocr(path: Path) -> str:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return ""
    result = subprocess.run(
        [tesseract, str(path), "stdout", "-l", "rus+eng"],
        check=False,
        capture_output=True,
        text=True,
        timeout=90,
    )
    return result.stdout if result.returncode == 0 else ""


def _extract_pdf_ocr(path: Path, max_pages: int) -> str:
    if max_pages <= 0 or not shutil.which("tesseract"):
        return ""
    try:
        import fitz

        doc = fitz.open(str(path))
        texts: list[str] = []
        with tempfile.TemporaryDirectory(prefix="aipoisk-pdf-ocr-") as tmp:
            for page_index in range(min(max_pages, len(doc))):
                page = doc[page_index]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_path = Path(tmp) / f"page-{page_index + 1}.png"
                pix.save(str(image_path))
                page_text = _extract_image_ocr(image_path)
                if page_text.strip():
                    texts.append(f"\n=== OCR PAGE {page_index + 1} ===\n{page_text}")
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_archive(path: Path, options: dict, depth: int) -> tuple[str, str]:
    max_files = int(options.get("archive_max_files") or 0)
    max_mb = int(options.get("archive_max_mb") or 0)
    if max_mb and path.stat().st_size > max_mb * 1024 * 1024:
        return "", "archive_too_large"
    with tempfile.TemporaryDirectory(prefix="aipoisk-archive-") as tmp:
        tmp_path = Path(tmp)
        if path.suffix.lower() == ".zip":
            extracted = _extract_zip_members(path, tmp_path, max_files)
        else:
            extracted = _extract_external_archive(path, tmp_path, max_files)
        if not extracted:
            return "", "archive_empty"
        parts: list[str] = []
        statuses: list[str] = []
        for item in extracted[: max_files or len(extracted)]:
            if item.is_dir():
                continue
            text, status = extract_text(item, options, _depth=depth - 1)
            statuses.append(status)
            if text.strip():
                rel = item.relative_to(tmp_path)
                parts.append(f"\n\n=== ARCHIVE FILE: {rel} ===\n{text}")
        if parts:
            status = "archive_ok" if all(item == "ok" for item in statuses) else "archive_partial"
            return "\n".join(parts), status
        return "", "archive_no_text"


def _extract_zip_members(path: Path, destination: Path, max_files: int) -> list[Path]:
    extracted: list[Path] = []
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist():
            if max_files and len(extracted) >= max_files:
                break
            if member.is_dir():
                continue
            safe_name = sanitize_filename(Path(member.filename).name)
            if not safe_name:
                continue
            target = destination / f"{len(extracted) + 1:03d}_{safe_name}"
            target.write_bytes(archive.read(member))
            extracted.append(target)
    return extracted


def _extract_external_archive(path: Path, destination: Path, max_files: int) -> list[Path]:
    if path.suffix.lower() == ".rar" and shutil.which("unrar"):
        command = ["unrar", "x", "-o+", str(path), str(destination)]
    elif shutil.which("7z"):
        command = ["7z", "x", f"-o{destination}", "-y", str(path)]
    else:
        return []
    result = subprocess.run(command, check=False, capture_output=True, text=True, timeout=180)
    if result.returncode != 0:
        return []
    files = [item for item in destination.rglob("*") if item.is_file()]
    return files[: max_files or len(files)]


def combined_document_context(items: list[tuple[str, str]]) -> str:
    parts = []
    for filename, text in items:
        if text.strip():
            parts.append(f"\n\n=== FILE: {filename} ===\n{text[:500000]}")
    return "\n".join(parts).strip()


def read_json_file(path: str | Path, default):
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except Exception:
        return default
